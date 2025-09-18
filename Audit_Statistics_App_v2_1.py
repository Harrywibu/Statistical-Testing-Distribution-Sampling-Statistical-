from __future__ import annotations
import os, io, re, json, time, hashlib, contextlib, tempfile, warnings
from datetime import datetime
from typing import Optional, List, Callable, Dict, Any
import numpy as np
import pandas as pd
import streamlit as st

def require_full_data(banner='Chưa có dữ liệu FULL. Hãy dùng **Load full data** trước khi chạy tab này.'):
    df = SS.get('df')
    import pandas as pd
    if df is None or not isinstance(df, pd.DataFrame) or df.empty:
        st.info(banner); st.stop()
    return df
    
def _first_match(colnames, patterns):
    cols = [c for c in colnames]
    low = {c: str(c).lower() for c in cols}
    for p in patterns:
        p = p.lower()
        for c in cols:
            if p in low[c]:
                return c
    return None

def _decode_bytes_to_str(v):
    if isinstance(v, (bytes, bytearray)):
        for enc in ('utf-8','latin-1','cp1252'):
            try:
                return v.decode(enc, errors='ignore')
            except Exception:
                pass
        try:
            return v.hex()
        except Exception:
            return str(v)
    return v

def sanitize_for_arrow(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or not isinstance(df, pd.DataFrame):
        return df
    df = df.copy()
    obj_cols = df.select_dtypes(include=['object']).columns
    for c in obj_cols:
        col = df[c]
        if col.isna().all():
            continue
        # bytes -> str
        if col.map(lambda v: isinstance(v, (bytes, bytearray))).any():
            df[c] = col.map(_decode_bytes_to_str)
            col = df[c]
        try:
            sample = col.dropna().iloc[:1000]
        except Exception:
            sample = col.dropna()
        has_str = any(isinstance(x, str) for x in sample)
        has_num = any(isinstance(x, (int,float,np.integer,np.floating)) for x in sample)
        has_nested = any(isinstance(x, (dict,list,set,tuple)) for x in sample)
        if has_nested or (has_str and has_num):
            df[c] = col.astype(str)
    return df

# ---- Streamlit width compatibility wrappers ----
try:
    _df_params = signature(st.dataframe).parameters
    _df_supports_width = 'width' in _df_params
except Exception:
    _df_supports_width = False



def st_df(data=None, **kwargs):
    if _df_supports_width:
        if kwargs.pop('use_container_width', None) is True:
            kwargs['width'] = 'stretch'
        elif 'width' not in kwargs:
            kwargs['width'] = 'stretch'
    else:
        kwargs.setdefault('use_container_width', True)
    return st.dataframe(data, **kwargs)  # Không gọi lại st_df

from scipy import stats

warnings.filterwarnings('ignore')

# Optional deps
try:
    import plotly.express as px
    import plotly.graph_objects as go
    from plotly.subplots import make_subplots
    HAS_PLOTLY = True
except Exception:
    HAS_PLOTLY = False
try:
    import plotly.io as pio
    HAS_KALEIDO = True
except Exception:
    HAS_KALEIDO = False
try:
    import docx
    from docx.shared import Inches
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False
try:
    import fitz  # PyMuPDF
    HAS_PDF = True
except Exception:
    HAS_PDF = False
try:
    import pyarrow as pa
    import pyarrow.parquet as pq
    HAS_PYARROW = True
except Exception:
    HAS_PYARROW = False
try:
    from sklearn.model_selection import train_test_split
    from sklearn.linear_model import LinearRegression, LogisticRegression
    from sklearn.metrics import (
        r2_score, mean_squared_error, accuracy_score, roc_auc_score, roc_curve
    )
    HAS_SK = True
except Exception:
    HAS_SK = False

# --------------------------------- App Config ---------------------------------
st.set_page_config(page_title='Audit Statistics', layout='wide', initial_sidebar_state='expanded')
SS = st.session_state

DEFAULTS = {
    'bins': 50,
    'log_scale': False,
    'kde_threshold': 150_000,
    'risk_diff_threshold': 0.05,
    'advanced_visuals': False,
    'use_parquet_cache': False,
    'pv_n': 100,
    'df': None,
 'last_good_df': None,
    'df_preview': None,
 'last_good_preview': None,
    'file_bytes': None,
 'ingest_ready': False,
    'sha12': '',
    'uploaded_name': '',
        'xlsx_sheet': '',
    'header_row': 1,
    'skip_top': 0,
 'col_whitelist': None,
}
for k, v in DEFAULTS.items():
    SS.setdefault(k, v)

# ------------------------------- Small Utilities ------------------------------
def file_sha12(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()[:12]

def st_plotly(fig, **kwargs):
    if '_plt_seq' not in SS: SS['_plt_seq'] = 0
    SS['_plt_seq'] += 1
    kwargs.setdefault('use_container_width', True)
    kwargs.setdefault('config', {'displaylogo': False})
    kwargs.setdefault('key', f'plt_{SS["_plt_seq"]}')
    return st.plotly_chart(fig, **kwargs)

@st.cache_data(ttl=900, show_spinner=False, max_entries=64)
def corr_cached(df: pd.DataFrame, cols: List[str], method: str = 'pearson') -> pd.DataFrame:
    if not cols: return pd.DataFrame()
    sub = df[cols].apply(pd.to_numeric, errors='coerce')
    sub = sub.dropna(axis=1, how='all')
    nunique = sub.nunique(dropna=True)
    keep = [c for c in sub.columns if nunique.get(c, 0) > 1]
    sub = sub[keep]
    if sub.shape[1] < 2: return pd.DataFrame()
    return sub.corr(method=method)

def is_datetime_like(colname: str, s: pd.Series) -> bool:
    return pd.api.types.is_datetime64_any_dtype(s) or bool(re.search(r'(date|time)', str(colname), re.I))

def _downcast_numeric(df: pd.DataFrame) -> pd.DataFrame:
    for c in df.select_dtypes(include=['float64']).columns:
        df[c] = pd.to_numeric(df[c], downcast='float')
    for c in df.select_dtypes(include=['int64']).columns:
        df[c] = pd.to_numeric(df[c], downcast='integer')
    return df

def to_float(x) -> Optional[float]:
    from numbers import Real
    try:
        if isinstance(x, Real): return float(x)
        if x is None: return None
        return float(str(x).strip().replace(',', ''))
    except Exception:
        return None

# ------------------------------- Disk Cache I/O --------------------------------
def _parquet_cache_path(sha: str, key: str) -> str:
    return os.path.join(tempfile.gettempdir(), f'astats_cache_{sha}_{key}.parquet')

@st.cache_data(ttl=6*3600, show_spinner=False, max_entries=24)
def write_parquet_cache(df: pd.DataFrame, sha: str, key: str) -> str:
    if not HAS_PYARROW: return ''
    try:
        df = sanitize_for_arrow(df)
        table = pa.Table.from_pandas(df)
        path = _parquet_cache_path(sha, key)
        pq.write_table(table, path)
        return path
    except Exception:
        return ''

def read_parquet_cache(sha: str, key: str) -> Optional[pd.DataFrame]:
    if not HAS_PYARROW: return None
    path = _parquet_cache_path(sha, key)
    if os.path.exists(path):
        try:
            return pq.read_table(path).to_pandas()
        except Exception:
            return None
    return None

# ------------------------------- Fast Readers ---------------------------------
@st.cache_data(ttl=6*3600, show_spinner=False, max_entries=16)
def list_sheets_xlsx(file_bytes: bytes) -> List[str]:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    try:
        return wb.sheetnames
    finally:
        wb.close()

@st.cache_data(ttl=6*3600, show_spinner=False, max_entries=16)
def read_csv_fast(file_bytes: bytes, usecols=None) -> pd.DataFrame:
    bio = io.BytesIO(file_bytes)
    try:
        df = pd.read_csv(bio, usecols=usecols, engine='pyarrow')
    except Exception:
        bio.seek(0)
        df = pd.read_csv(bio, usecols=usecols, low_memory=False, memory_map=True)
    return _downcast_numeric(df)

@st.cache_data(ttl=6*3600, show_spinner=False, max_entries=16)
def read_xlsx_fast(file_bytes: bytes, sheet: str, usecols=None, header_row: int = 1, skip_top: int = 0, dtype_map=None) -> pd.DataFrame:
    skiprows = list(range(header_row, header_row + skip_top)) if skip_top > 0 else None
    bio = io.BytesIO(file_bytes)
    df = pd.read_excel(bio, sheet_name=sheet, usecols=usecols, header=header_row - 1,
                       skiprows=skiprows, dtype=dtype_map, engine='openpyxl')
    return _downcast_numeric(df)

# ----------------------------- Cached Basic Stats -----------------------------
@st.cache_data(ttl=1800, show_spinner=False, max_entries=64)
def numeric_profile_stats(series: pd.Series):
    s = pd.to_numeric(series, errors='coerce').replace([np.inf, -np.inf], np.nan).dropna()
    desc = s.describe(percentiles=[0.01,0.05,0.10,0.25,0.5,0.75,0.90,0.95,0.99])
    skew = float(stats.skew(s)) if len(s) > 2 else np.nan
    kurt = float(stats.kurtosis(s, fisher=True)) if len(s) > 3 else np.nan
    try:
        p_norm = float(stats.normaltest(s)[1]) if len(s) > 7 else np.nan
    except Exception:
        p_norm = np.nan
    p95 = s.quantile(0.95) if len(s)>1 else np.nan
    p99 = s.quantile(0.99) if len(s)>1 else np.nan
    zero_ratio = float((s==0).mean()) if len(s)>0 else np.nan
    return desc.to_dict(), skew, kurt, p_norm, float(p95), float(p99), zero_ratio

@st.cache_data(ttl=1800, show_spinner=False, max_entries=64)
def cat_freq(series: pd.Series) -> pd.DataFrame:
    s = series.dropna().astype(str)
    vc = s.value_counts(dropna=True)
    out = pd.DataFrame({'category': vc.index, 'count': vc.values})
    out['share'] = out['count']/out['count'].sum()
    return out


# ------------------------------ Benford Helpers -------------------------------
@st.cache_data(ttl=3600, show_spinner=False, max_entries=64)
def _benford_1d(series: pd.Series):
    s = pd.to_numeric(series, errors='coerce').replace([np.inf, -np.inf], np.nan).dropna().abs()
    if s.empty: return None
    def _digits(x):
        xs = ("%.15g" % float(x))
        return re.sub(r"[^0-9]", "", xs).lstrip("0")
    d1 = s.apply(lambda v: int(_digits(v)[0]) if len(_digits(v))>=1 else np.nan).dropna()
    d1 = d1[(d1>=1)&(d1<=9)]
    if d1.empty: return None
    obs = d1.value_counts().sort_index().reindex(range(1,10), fill_value=0).astype(float)
    n=obs.sum(); obs_p=obs/n
    idx=np.arange(1,10); exp_p=np.log10(1+1/idx); exp=exp_p*n
    with np.errstate(divide='ignore', invalid='ignore'):
        chi2=np.nansum((obs-exp)**2/exp)
        pval=1-stats.chi2.cdf(chi2, len(idx)-1)
        mad=float(np.mean(np.abs(obs_p-exp_p)))
        var_tbl=pd.DataFrame({'digit':idx,'expected':exp,'observed':obs.values})
        var_tbl['diff']=var_tbl['observed']-var_tbl['expected']
        var_tbl['diff_pct']=(var_tbl['observed']-var_tbl['expected'])/var_tbl['expected']
        table=pd.DataFrame({'digit':idx,'observed_p':obs_p.values,'expected_p':exp_p})
    return {'table':table, 'variance':var_tbl, 'n':int(n), 'chi2':float(chi2), 'p':float(pval), 'MAD':float(mad)}

@st.cache_data(ttl=3600, show_spinner=False, max_entries=64)
def _benford_2d(series: pd.Series):
    s = pd.to_numeric(series, errors='coerce').replace([np.inf, -np.inf], np.nan).dropna().abs()
    if s.empty: return None
    def _digits(x):
        xs = ("%.15g" % float(x))
        return re.sub(r"[^0-9]", "", xs).lstrip("0")
    def _first2(v):
        ds = _digits(v)
        if len(ds)>=2: return int(ds[:2])
        if len(ds)==1 and ds!="0": return int(ds)
        return np.nan
    d2 = s.apply(_first2).dropna(); d2=d2[(d2>=10)&(d2<=99)]
    if d2.empty: return None
    obs = d2.value_counts().sort_index().reindex(range(10,100), fill_value=0).astype(float)
    n=obs.sum(); obs_p=obs/n
    idx=np.arange(10,100); exp_p=np.log10(1+1/idx); exp=exp_p*n
    with np.errstate(divide='ignore', invalid='ignore'):
        chi2=np.nansum((obs-exp)**2/exp)
        pval=1-stats.chi2.cdf(chi2, len(idx)-1)
    mad=float(np.mean(np.abs(obs_p-exp_p)))
    var_tbl=pd.DataFrame({'digit':idx,'expected':exp,'observed':obs.values})
    var_tbl['diff']=var_tbl['observed']-var_tbl['expected']
    var_tbl['diff_pct']=(var_tbl['observed']-var_tbl['expected'])/var_tbl['expected']
    table=pd.DataFrame({'digit':idx,'observed_p':obs_p.values,'expected_p':exp_p})
    return {'table':table, 'variance':var_tbl, 'n':int(n), 'chi2':float(chi2), 'p':float(pval), 'MAD':float(mad)}

def _benford_ready(series: pd.Series) -> tuple[bool, str]:
    s = pd.to_numeric(series, errors='coerce')
    n_nz = int((s != 0).sum())  # nhận cả số âm, chỉ loại 0
    if n_nz < 1:
        return False, f"Không có giá trị ≠ 0 để chạy Benford (hiện {n_nz}, cần ≥300)."
    s_non = s.dropna()
    if s_non.shape[0] > 0:
        ratio_unique = s_non.nunique()/s_non.shape[0]
        if ratio_unique > 0.95:
            return False, "Tỉ lệ unique quá cao (khả năng ID/Code) — tránh Benford."
    return True, ''

def _plot(fig):
    try:
        st_plotly(fig)
    except Exception:
        st.plotly_chart(fig, use_container_width=True)

def guess_datetime_cols(df, check=3000):
    import numpy as np, pandas as pd
    sample = df.head(check)
    cols = []
    for c in df.columns:
        try:
            if np.issubdtype(df[c].dtype, np.datetime64):
                cols.append(c); continue
            if df[c].dtype == 'object':
                s = pd.to_datetime(sample[c], errors='coerce')
                if s.notna().mean() >= 0.5:
                    cols.append(c)
        except Exception:
            pass
    return cols

# -------------------------- Sidebar: Workflow & perf ---------------------------
up = st.file_uploader('Upload file (.csv, .xlsx)', type=['csv','xlsx'], key='ingest')
if up is not None:
    fb = up.read()  # có thể dùng up.getvalue() cũng được
    new_sha = file_sha12(fb)
    same_file = (SS.get('sha12') == new_sha) and (SS.get('uploaded_name') == up.name)

    # luôn cập nhật metadata/bytes để các bước sau dùng
    SS['file_bytes'] = fb
    SS['uploaded_name'] = up.name
    SS['sha12'] = new_sha

    # 🔒 CHỈ khi đổi file mới reset preview/full
    if not same_file:
        SS['df'] = None
        SS['df_preview'] = None

    st.caption(f"Đã nhận file: {up.name} • SHA12={SS['sha12']}")

    if st.button('Clear file', key='btn_clear_file'):
        base_keys = ['file_bytes','uploaded_name','sha12','df','df_preview','col_whitelist']
        result_keys = [
            'bf1_res','bf2_res','bf1_col','bf2_col','t4_results','last_corr','last_linear',
            'last_logistic','last_numeric_profile','last_gof','fraud_flags','spearman_recommended',
            '_plt_seq','col_filter','dtype_choice','xlsx_sheet','header_row','skip_top',
            'ingest_ready','last_good_df','last_good_preview'
        ]
        # đặt tên biến khác nhau để tránh đè 'k'
        for bk in base_keys:
            SS[bk] = DEFAULTS.get(bk, None)

        for rk in result_keys:
            if rk in SS:
                SS[rk] = None

        st.rerun()
with st.sidebar.expander('1) Display & Performance', expanded=True):
    SS['bins'] = st.slider('Histogram bins', 10, 200, SS.get('bins',50), 5)
    SS['log_scale'] = st.checkbox('Log scale (X)', value=SS.get('log_scale', False))
    SS['kde_threshold'] = st.number_input('KDE max n', 1_000, 300_000, SS.get('kde_threshold',150_000), 1_000)
with st.sidebar.expander('2) Risk & Advanced', expanded=False):
    SS['risk_diff_threshold'] = st.slider('Benford diff% threshold', 0.01, 0.10, SS.get('risk_diff_threshold',0.05), 0.01)
    SS['advanced_visuals'] = st.checkbox('Advanced visuals (Violin, Lorenz/Gini)', value=SS.get('advanced_visuals', False))
with st.sidebar.expander('3) Cache', expanded=False):
    if not HAS_PYARROW:
        st.caption('⚠️ PyArrow chưa sẵn sàng — Disk cache (Parquet) sẽ bị tắt.')
        SS['use_parquet_cache'] = False
    SS['use_parquet_cache'] = st.checkbox('Disk cache (Parquet) for faster reloads', value=SS.get('use_parquet_cache', False) and HAS_PYARROW)
    if st.button('🧹 Clear cache'):
        st.cache_data.clear(); st.toast('Cache cleared', icon='🧹')

# ---------------------------------- Main Gate ---------------------------------
st.title('📊 Audit Statistics')
if SS['file_bytes'] is None:
    st.info('Upload a file để bắt đầu.'); st.stop()

fname=SS['uploaded_name']; fb=SS['file_bytes']; sha=SS['sha12']
colL, colR = st.columns([3,2])
with colL:
    st.text_input('File', value=fname or '', disabled=True)
with colR:
    SS['pv_n'] = st.slider('Preview rows', 50, 500, SS.get('pv_n',100), 50)
    do_preview = st.button('🔎 Quick preview', key='btn_prev')

# Ingest flow
if fname.lower().endswith('.csv'):
    if do_preview or SS['df_preview'] is None:
        try:
            SS['df_preview'] = sanitize_for_arrow(read_csv_fast(fb).head(SS['pv_n']))
            SS['last_good_preview'] = SS['df_preview']; SS['ingest_ready']=True
        except Exception as e:
            st.error(f'Lỗi đọc CSV: {e}'); SS['df_preview']=None
    if SS['df_preview'] is not None:
        st_df(SS['df_preview'], use_container_width=True, height=260)
        headers=list(SS['df_preview'].columns)
        selected = st.multiselect('Columns to load', headers, default=headers)
        SS['col_whitelist'] = selected if selected else headers
        if st.button('📥 Load full CSV with selected columns', key='btn_load_csv'):
            sel_key=';'.join(selected) if selected else 'ALL'
            key=f"csv_{hashlib.sha1(sel_key.encode()).hexdigest()[:10]}"
            df_cached = read_parquet_cache(sha, key) if SS['use_parquet_cache'] else None
            if df_cached is None:
                df_full = sanitize_for_arrow(read_csv_fast(fb, usecols=(selected or None)))
                if SS['use_parquet_cache']: write_parquet_cache(df_full, sha, key)
            else:
                df_full = df_cached
            SS['df']=df_full; SS['last_good_df']=df_full; SS['ingest_ready']=True; SS['col_whitelist']=list(df_full.columns)
            st.success(f"Loaded: {len(SS['df']):,} rows × {len(SS['df'].columns)} cols • SHA12={sha}")
else:
    sheets = list_sheets_xlsx(fb)
    with st.expander('📁 Select sheet & header (XLSX)', expanded=True):
        c1,c2,c3 = st.columns([2,1,1])
        idx=0 if sheets else 0
        SS['xlsx_sheet'] = c1.selectbox('Sheet', sheets, index=idx)
        SS['header_row'] = c2.number_input('Header row (1‑based)', 1, 100, SS['header_row'])
        SS['skip_top'] = c3.number_input('Skip N rows after header', 0, 1000, SS['skip_top'])
        SS['dtype_choice'] = st.text_area('dtype mapping (JSON, optional)', SS.get('dtype_choice',''), height=60)
        dtype_map=None
        if SS['dtype_choice'].strip():
            try: dtype_map=json.loads(SS['dtype_choice'])
            except Exception as e: st.warning(f'Không đọc được dtype JSON: {e}')
        try:
            prev = sanitize_for_arrow(read_xlsx_fast(fb, SS['xlsx_sheet'], usecols=None, header_row=SS['header_row'], skip_top=SS['skip_top'], dtype_map=dtype_map).head(SS['pv_n']))
            SS['df_preview']=prev; SS['last_good_preview']=prev; SS['ingest_ready']=True
        except Exception as e:
            st.error(f'Lỗi đọc XLSX: {e}'); prev=pd.DataFrame()
        st_df(prev, use_container_width=True, height=260)
        headers=list(prev.columns)
        st.caption(f'Columns: {len(headers)} • SHA12={sha}')
        SS['col_filter'] = st.text_input('🔎 Filter columns', SS.get('col_filter',''))
        filtered = [h for h in headers if SS['col_filter'].lower() in h.lower()] if SS['col_filter'] else headers
        selected = st.multiselect('🧮 Columns to load', filtered if filtered else headers, default=filtered if filtered else headers)
        if st.button('📥 Load full data', key='btn_load_xlsx'):
            key_tuple=(SS['xlsx_sheet'], SS['header_row'], SS['skip_top'], tuple(selected) if selected else ('ALL',))
            key=f"xlsx_{hashlib.sha1(str(key_tuple).encode()).hexdigest()[:10]}"
            df_cached = read_parquet_cache(sha, key) if SS['use_parquet_cache'] else None
            if df_cached is None:
                df_full = sanitize_for_arrow(read_xlsx_fast(fb, SS['xlsx_sheet'], usecols=(selected or None), header_row=SS['header_row'], skip_top=SS['skip_top'], dtype_map=dtype_map))
                if SS['use_parquet_cache']: write_parquet_cache(df_full, sha, key)
            else:
                df_full = df_cached
            SS['df']=df_full; SS['last_good_df']=df_full; SS['ingest_ready']=True; SS['col_whitelist']=list(df_full.columns)
            st.success(f"Loaded: {len(SS['df']):,} rows × {len(SS['df'].columns)} cols • SHA12={sha}")

if SS['df'] is None and SS['df_preview'] is None:
    st.stop()

# Source & typing
DF_FULL = require_full_data('Chưa có dữ liệu FULL. Hãy dùng **Load full data**.')
DF_VIEW = DF_FULL  # alias để không phá code cũ

ALL_COLS = list(DF_FULL.columns)
DT_COLS  = [c for c in ALL_COLS if is_datetime_like(c, DF_FULL[c])]
NUM_COLS = DF_FULL[ALL_COLS].select_dtypes(include=[np.number]).columns.tolist()
CAT_COLS = DF_FULL[ALL_COLS].select_dtypes(include=['object','category','bool']).columns.tolist()
VIEW_COLS = [c for c in DF_FULL.columns if (not SS.get('col_whitelist') or c in SS['col_whitelist'])]


@st.cache_data(ttl=900, show_spinner=False, max_entries=64)
def spearman_flag(df: pd.DataFrame, cols: List[str]) -> bool:
    try:
        if df is None or not isinstance(df, pd.DataFrame):
            return False
    except Exception:
        return False

    for c in (cols or [])[:20]:
        if c not in df.columns:
            continue

        s = pd.to_numeric(df[c], errors='coerce').replace([np.inf, -np.inf], np.nan).dropna()
        if len(s) < 50:
            continue

        sk, ku, tail, p_norm = 0.0, 0.0, 0.0, 1.0  # defaults

        try:
            if len(s) > 2:
                sk = float(stats.skew(s))
        except Exception:
            pass

        try:
            if len(s) > 3:
                ku = float(stats.kurtosis(s, fisher=True))
        except Exception:
            pass

        try:
            p99 = s.quantile(0.99)
            if pd.notna(p99):
                tail = float((s > p99).mean())
        except Exception:
            pass

        try:
            if len(s) > 20:
                p_norm = float(stats.normaltest(s)[1])
        except Exception:
            pass

        if (abs(sk) > 1) or (abs(ku) > 3) or (tail > 0.02) or (p_norm < 0.05):
            return True
    return False
# ------------------------------ Rule Engine Core ------------------------------
class Rule:
    def __init__(self, id: str, name: str, scope: str, severity: str,
                 condition: Callable[[Dict[str,Any]], bool],
                 action: str, rationale: str):
        self.id=id; self.name=name; self.scope=scope; self.severity=severity
        self.condition=condition; self.action=action; self.rationale=rationale

    def eval(self, ctx: Dict[str,Any]) -> Optional[Dict[str,Any]]:
        try:
            if self.condition(ctx):
                return {
                    'id': self.id,
                    'name': self.name,
                    'scope': self.scope,
                    'severity': self.severity,
                    'action': self.action,
                    'rationale': self.rationale,
                }
        except Exception:
            return None
        return None

SEV_ORDER = {'High':3,'Medium':2,'Low':1,'Info':0}

def _get(ctx: Dict[str,Any], *keys, default=None):
    cur = ctx
    for k in keys:
        if cur is None: return default
        cur = cur.get(k) if isinstance(cur, dict) else None
    return cur if cur is not None else default

def build_rule_context() -> Dict[str,Any]:
    ctx = {
        'thr': {
            'benford_diff': SS.get('risk_diff_threshold', 0.05),
            'zero_ratio': 0.30,
            'tail_p99': 0.02,
            'off_hours': 0.15,
            'weekend': 0.25,
            'corr_high': 0.9,
            'gap_p95_hours': 12.0,
            'hhi': 0.2,
        },
        'last_numeric': SS.get('last_numeric_profile'),
        'gof': SS.get('last_gof'),
        'benford': {
            'r1': SS.get('bf1_res'),
            'r2': SS.get('bf2_res')
        },
        't4': SS.get('t4_results'),
        'corr': SS.get('last_corr'),
        'regression': {
            'linear': SS.get('last_linear'),
            'logistic': SS.get('last_logistic'),
        },
        'flags': SS.get('fraud_flags') or [],
    }
    # convenience derivations
    r1 = ctx['benford'].get('r1')
    if r1 and isinstance(r1, dict) and 'variance' in r1:
        try:
            ctx['benford']['r1_maxdiff'] = float(r1['variance']['diff_pct'].abs().max())
        except Exception:
            ctx['benford']['r1_maxdiff'] = None
    r2 = ctx['benford'].get('r2')
    if r2 and isinstance(r2, dict) and 'variance' in r2:
        try:
            ctx['benford']['r2_maxdiff'] = float(r2['variance']['diff_pct'].abs().max())
        except Exception:
            ctx['benford']['r2_maxdiff'] = None
    return ctx

def rules_catalog() -> List[Rule]:
    R: List[Rule] = []
    # Profiling — zero heavy
    R.append(Rule(
        id='NUM_ZERO_HEAVY', name='Zero‑heavy numeric', scope='profiling', severity='Medium',
        condition=lambda c: _get(c,'last_numeric','zero_ratio', default=0)<=1 and _get(c,'last_numeric','zero_ratio', default=0) > _get(c,'thr','zero_ratio'),
        action='Kiểm tra policy/threshold; χ² tỷ lệ theo đơn vị/nhóm; cân nhắc data quality.',
        rationale='Tỉ lệ 0 cao có thể do ngưỡng phê duyệt/không sử dụng trường/ETL.'
    ))
    # Profiling — heavy right tail
    R.append(Rule(
        id='NUM_TAIL_HEAVY', name='Đuôi phải dày (>P99)', scope='profiling', severity='High',
        condition=lambda c: _get(c,'last_numeric','tail_gt_p99', default=0) > _get(c,'thr','tail_p99'),
        action='Benford 1D/2D; xem cut‑off cuối kỳ; rà soát outliers/drill‑down.',
        rationale='Đuôi phải dày liên quan bất thường giá trị lớn/outliers.'
    ))
    # GoF suggests transform
    R.append(Rule(
        id='GOF_TRANSFORM', name='Nên biến đổi (log/Box‑Cox)', scope='profiling', severity='Info',
        condition=lambda c: bool(_get(c,'gof','suggest')) and _get(c,'gof','best') in {'Lognormal','Gamma'},
        action='Áp dụng log/Box‑Cox trước các test tham số hoặc dùng phi tham số.',
        rationale='Phân phối lệch/không chuẩn — biến đổi giúp thỏa giả định tham số.'
    ))
    # Benford 1D
    R.append(Rule(
        id='BENFORD_1D_SEV', name='Benford 1D lệch', scope='benford', severity='High',
        condition=lambda c: (_get(c,'benford','r1') is not None) and \
            ((_get(c,'benford','r1','p', default=1.0) < 0.05) or (_get(c,'benford','r1','MAD', default=0) > 0.012) or \
             (_get(c,'benford','r1_maxdiff', default=0) >= _get(c,'thr','benford_diff'))),
        action='Drill‑down nhóm digit chênh nhiều; đối chiếu nhà CC/kỳ; kiểm tra cut‑off.',
        rationale='Lệch Benford gợi ý thresholding/làm tròn/chia nhỏ hóa đơn.'
    ))
    # Benford 2D
    R.append(Rule(
        id='BENFORD_2D_SEV', name='Benford 2D lệch', scope='benford', severity='Medium',
        condition=lambda c: (_get(c,'benford','r2') is not None) and \
            ((_get(c,'benford','r2','p', default=1.0) < 0.05) or (_get(c,'benford','r2','MAD', default=0) > 0.012) or \
             (_get(c,'benford','r2_maxdiff', default=0) >= _get(c,'thr','benford_diff'))),
        action='Xem hot‑pair (19/29/…); đối chiếu chính sách giá; không mặc định là gian lận.',
        rationale='Mẫu cặp chữ số đầu bất thường có thể phản ánh hành vi định giá.'
    ))
    # Categorical — HHI high
    R.append(Rule(
        id='HHI_HIGH', name='Tập trung nhóm cao (HHI)', scope='tests', severity='Medium',
        condition=lambda c: _get(c,'t4','hhi','hhi', default=0) > _get(c,'thr','hhi'),
        action='Đánh giá rủi ro phụ thuộc nhà cung cấp/GL; kiểm soát phê duyệt.',
        rationale='HHI cao cho thấy rủi ro tập trung vào ít nhóm.'
    ))
    # Categorical — Chi-square significant
    R.append(Rule(
        id='CGOF_SIG', name='Chi‑square GoF khác Uniform', scope='tests', severity='Medium',
        condition=lambda c: _get(c,'t4','cgof','p', default=1.0) < 0.05,
        action='Drill‑down residual lớn; xem data quality/policy phân loại.',
        rationale='Sai khác mạnh so với uniform gợi ý phân phối lệch có chủ đích.'
    ))
    # Time — Gap large
    R.append(Rule(
        id='TIME_GAP_LARGE', name='Khoảng cách thời gian lớn (p95)', scope='tests', severity='Low',
        condition=lambda c: to_float(_get(c,'t4','gap','gaps','gap_hours','describe','95%', default=np.nan)) or False,
        action='Xem kịch bản bỏ sót/chèn nghiệp vụ; đối chiếu lịch chốt.',
        rationale='Khoảng trống dài bất thường có thể do quy trình/ghi nhận không liên tục.'
    ))
    # Correlation — high multicollinearity
    def _corr_high(c: Dict[str,Any]):
        M = _get(c,'corr');
        if not isinstance(M, pd.DataFrame) or M.empty: return False
        thr = _get(c,'thr','corr_high', default=0.9)
        tri = M.where(~np.eye(len(M), dtype=bool))
        return np.nanmax(np.abs(tri.values)) >= thr
    R.append(Rule(
        id='CORR_HIGH', name='Tương quan rất cao giữa biến', scope='correlation', severity='Info',
        condition=_corr_high,
        action='Kiểm tra đa cộng tuyến; cân nhắc loại bớt biến khi hồi quy.',
        rationale='|r| cao gây bất ổn ước lượng tham số.'
    ))
    # Flags — duplicates
    def _flags_dup(c: Dict[str,Any]):
        return any((isinstance(x, dict) and 'Duplicate' in str(x.get('flag',''))) for x in _get(c,'flags', default=[]))
    R.append(Rule(
        id='DUP_KEYS', name='Trùng khóa/tổ hợp', scope='flags', severity='High',
        condition=_flags_dup,
        action='Rà soát entries trùng; kiểm soát nhập liệu/phê duyệt; root‑cause.',
        rationale='Trùng lặp có thể là double posting/ghost entries.'
    ))
    # Flags — off hours/weekend
    def _flags_off(c):
        return any('off-hours' in str(x.get('flag','')).lower() for x in _get(c,'flags', default=[]))
    R.append(Rule(
        id='OFF_HOURS', name='Hoạt động off‑hours/ cuối tuần', scope='flags', severity='Medium',
        condition=_flags_off,
        action='Rà soát phân quyền/ca trực/automation; χ² theo khung giờ × status.',
        rationale='Hoạt động bất thường ngoài giờ có thể là tín hiệu rủi ro.'
    ))
    # Regression — poor linear fit
    R.append(Rule(
        id='LIN_POOR', name='Linear Regression kém (R2 thấp)', scope='regression', severity='Info',
        condition=lambda c: to_float(_get(c,'regression','linear','R2')) is not None and to_float(_get(c,'regression','linear','R2')) < 0.3,
        action='Xem lại chọn biến/biến đổi/log/phi tuyến hoặc dùng mô hình khác.',
        rationale='R2 thấp: mô hình chưa giải thích tốt biến thiên mục tiêu.'
    ))
    # Regression — logistic good AUC
    R.append(Rule(
        id='LOGIT_GOOD', name='Logistic phân biệt tốt (AUC ≥ 0.7)', scope='regression', severity='Info',
        condition=lambda c: to_float(_get(c,'regression','logistic','ROC_AUC')) is not None and to_float(_get(c,'regression','logistic','ROC_AUC')) >= 0.7,
        action='Dùng model hỗ trợ ưu tiên kiểm thử; xem fairness & leakage.',
        rationale='AUC cao: có cấu trúc dự đoán hữu ích cho điều tra rủi ro.'
    ))
    return R

def evaluate_rules(ctx: Dict[str,Any], scope: Optional[str]=None) -> pd.DataFrame:
    rows=[]
    for r in rules_catalog():
        if scope and r.scope!=scope: continue
        hit = r.eval(ctx)
        if hit: rows.append(hit)
    if not rows: return pd.DataFrame(columns=['severity','scope','name','action','rationale'])
    df = pd.DataFrame(rows)
    df['sev_rank'] = df['severity'].map(SEV_ORDER).fillna(0)
    df = df.sort_values(['sev_rank','scope','name'], ascending=[False, True, True]).drop(columns=['sev_rank'])
    return df

# ----------------------------------- TABS -------------------------------------
TAB0, TAB1, TAB2, TAB3, TAB4, TAB5, TAB6, TAB7 = st.tabs([ '0) Data Quality (FULL)', '1) Overview (Sales activity)', '2) Profiling/Distribution', '3) Correlation & Trend', '4) Benford', '5) ANOVA & Nonparametric', '6) Regression', '7) Flags & Risk/Export'])
# ---- TAB 0: Data Quality (FULL) ----
with TAB0:
    st.subheader('🧪 Data Quality')
    if SS.get('df') is None:
        st.info('Hãy **Load full data** để xem Data Quality Tab.')
    else:
        @st.cache_data(ttl=900, show_spinner=False, max_entries=16)
        def data_quality_table(df_in):
            import pandas as pd
            rows = []
            n = len(df_in)
            for c in df_in.columns:
                s = df_in[c]
                is_num = pd.api.types.is_numeric_dtype(s)
                is_dt  = pd.api.types.is_datetime64_any_dtype(s) or is_datetime_like(c, s)
                is_bool= pd.api.types.is_bool_dtype(s)
                is_cat = pd.api.types.is_categorical_dtype(s)
                base_type = 'Numeric' if is_num else ('Datetime' if is_dt else ('Boolean' if is_bool else ('Categorical' if is_cat else 'Text')))
                n_nonnull = int(s.notna().sum())
                n_nan = int(n - n_nonnull)
                n_unique = int(s.nunique(dropna=True))
                mem_mb = float(s.memory_usage(deep=True)) / 1048576.0
                blank = None; blank_pct = None
                zero = None; zero_pct = None
                if base_type in ('Text','Categorical'):
                    s_txt = s[s.notna()].astype(str).str.strip()
                    blank = int((s_txt == '').sum())
                    blank_pct = round(blank / n, 4) if n else None
                if base_type == 'Numeric':
                    s_num = pd.to_numeric(s, errors='coerce')
                    zero = int(s_num.eq(0).sum())
                    zero_pct = round(zero / n, 4) if n else None
                valid = n_nonnull - (blank or 0) if base_type in ('Text','Categorical') else n_nonnull
                rows.append({
                    'column': c,
                    'type': base_type,
                    'rows': n,
                    'valid': int(valid),
                    'valid%': round(valid / n, 4) if n else None,
                    'nan': n_nan,
                    'nan%': round(n_nan / n, 4) if n else None,
                    'blank': blank,
                    'blank%': blank_pct,
                    'zero': zero,
                    'zero%': zero_pct,
                    'unique': n_unique,
                    'memory_MB': round(mem_mb, 3),
                })
            cols_order = ['column','type','rows','valid','valid%','nan','nan%','blank','blank%','zero','zero%','unique','memory_MB']
            dq = pd.DataFrame(rows)
            dq = dq[cols_order]
            return dq.sort_values(['type','column']).reset_index(drop=True)
        try:
            dq = data_quality_table(SS['df'] if SS.get('df') is not None else DF_VIEW)
            st_df(dq, use_container_width=True, height=min(520, 60 + 24*min(len(dq), 18)))
        except Exception as e:
            st.error(f'Lỗi Data Quality: {e}')
# ================================= TAB 1 — OVERVIEW (Sales Activities) =================================
with TAB1:
    import pandas as pd, numpy as np
    import plotly.graph_objects as go
    import streamlit as st

    st.subheader("📈 Overview — Sales Activities")

    # ===== Guard =====
    df = SS.get("df")
    if df is None or df.empty:
        st.info("Hãy nạp dữ liệu trước.")
        st.stop()

    # ===== Helpers =====
    def _pick(col, label, key):
        val = col.selectbox(label, ["—"] + list(df.columns), index=0, key=key)
        return None if val == "—" else val

    def _norm_period_value(p):
        if p is None: return "Month"
        s = str(p).strip().lower()
        if s in {"m","mo","mon","month","tháng"}: return "Month"
        if s in {"q","quy","quarter","quý"}:     return "Quarter"
        if s in {"y","yr","year","năm"}:         return "Year"
        if "quý" in s or s.startswith("q"):      return "Quarter"
        if "năm" in s or s.startswith("y"):      return "Year"
        return "Month"

    RULE_MAP   = {"Month":"MS","Quarter":"QS","Year":"YS"}           # resample rule
    PERIOD_MAP = {"MS":"M","QS":"Q","YS":"Y"}                        # to_period code
    YOY_LAG    = {"MS":12,"QS":4,"YS":1}

  # ================================= TAB 1 — OVERVIEW (Sales Activities) =================================
with TAB1:
    import pandas as pd, numpy as np
    import plotly.graph_objects as go
    import streamlit as st

    st.subheader("📈 Overview — Sales Activities")

    # ===== Data guard =====
    df = SS.get("df")
    if df is None or df.empty:
        st.info("Hãy nạp dữ liệu trước.")
        st.stop()

    # ===== Helpers =====
    def _pick(col, label, key):
        val = col.selectbox(label, ["—"] + list(df.columns), index=0, key=key)
        return None if val == "—" else val

    def _norm_period_value(p):
        if p is None: return "Month"
        s = str(p).strip().lower()
        if s in {"m","mo","mon","month","tháng"}: return "Month"
        if s in {"q","quy","quarter","quý"}:     return "Quarter"
        if s in {"y","yr","year","năm"}:         return "Year"
        if "quý" in s or s.startswith("q"):      return "Quarter"
        if "năm" in s or s.startswith("y"):      return "Year"
        return "Month"

    RULE_MAP   = {"Month":"MS","Quarter":"QS","Year":"YS"}    # resample rule
    PERIOD_MAP = {"MS":"M","QS":"Q","YS":"Y"}                 # to_period code
    YOY_LAG    = {"MS":12,"QS":4,"YS":1}

    # ============================= 0) CẤU HÌNH DỮ LIỆU — 2 HÀNG =============================
    st.markdown("### ⚙️ Cấu hình dữ liệu (bắt buộc) — 2 hàng")
    with st.container(border=True):
        # HÀNG 1 — Time / IDs / Dimensions
        c1, c2, c3, c4, c5, c6 = st.columns([1,1,1,1,1,1])
        time_col    = _pick(c1, "🕒 Time",        "cfg_time")
        order_col   = _pick(c2, "🧾 Order/Doc",   "cfg_order")
        cust_col    = _pick(c3, "👤 Customer",    "cfg_cust")
        prod_col    = _pick(c4, "📦 Product",     "cfg_prod")
        region_col  = _pick(c5, "🌍 Region",      "cfg_region")
        channel_col = _pick(c6, "🛒 Channel",     "cfg_channel")

        # HÀNG 2 — Schema + Value columns (gọn, song song)
        left, right = st.columns([0.9, 3.1])
        schema = left.segmented_control(
            "Schema",
            ["Amount + Type (2-type cols)", "Separate numeric cols"],
            key="cfg_schema"
        )

        if schema == "Amount + Type (2-type cols)":
            a1, a2, a3 = right.columns([1,1,1])
            amt_col  = _pick(a1, "💰 Amount",   "cfg_amt")
            type_col = _pick(a2, "🏷️ Txn type", "cfg_txn_type")   # Sales / Purchase / Transfer-in / Transfer-out / Returns
            adj_col  = _pick(a3, "🏷️ Adj type", "cfg_adj_type")   # Sales / Discount

            uniq_txn = list(pd.Series(df[type_col].astype(str).unique()).sort_values())[:2000] if type_col else []
            uniq_adj = list(pd.Series(df[adj_col].astype(str).unique()).sort_values())[:2000] if adj_col else []

            with st.expander("Mapping Txn (Sales / Purchase / Transfer-in / Transfer-out / Returns)", expanded=False):
                t1, t2, t3, t4, t5 = st.columns(5)
                val_txn_sales = t1.multiselect("Sales",        uniq_txn, key="map_txn_sales")
                val_purchase  = t2.multiselect("Purchase",     uniq_txn, key="map_purchase")
                val_tin       = t3.multiselect("Transfer-in",  uniq_txn, key="map_tin")
                val_tout      = t4.multiselect("Transfer-out", uniq_txn, key="map_tout")
                val_returns   = t5.multiselect("Returns",      uniq_txn, key="map_returns")

            with st.expander("Mapping Adj (Sales / Discount)", expanded=False):
                a1_, a2_ = st.columns(2)
                val_adj_sales = a1_.multiselect("Adj = Sales",    uniq_adj, key="map_adj_sales")
                val_adj_disc  = a2_.multiselect("Adj = Discount", uniq_adj, key="map_adj_disc")

            # đảm bảo biến tồn tại nếu user không mở expander
            for _v in ["val_txn_sales","val_purchase","val_tin","val_tout","val_returns","val_adj_sales","val_adj_disc"]:
                if _v not in locals(): locals()[_v] = []

        else:
            b1, b2, b3, b4, b5 = right.columns([1,1,1,1,1])
            sales_col   = _pick(b1, "Sales",             "cfg_sales")
            returns_col = _pick(b2, "Returns (opt)",     "cfg_ret")
            disc_col    = _pick(b3, "Discount (opt)",    "cfg_disc")
            tin_col     = _pick(b4, "Transfer-in (opt)", "cfg_tin")
            tout_col    = _pick(b5, "Transfer-out (opt)","cfg_tout")
            # tạo biến rỗng cho nhánh còn lại
            val_txn_sales = val_purchase = val_tin = val_tout = val_returns = val_adj_sales = val_adj_disc = []
            type_col = adj_col = None

    # ============================= 1) CẤU HÌNH HIỂN THỊ (rút gọn) =============================
    st.markdown("### 🧭 Cấu hình hiển thị")
    c1, c2, c3, c4 = st.columns([1,1,1.6,1.1])
    period_raw = c1.segmented_control("Period", ["Month","Quarter","Year"])
    compare    = c2.segmented_control("Compare", ["Prev","YoY"])
    dim_col    = c3.selectbox("📊 Dimension (X cho 'Đóng góp')", ["—"] + list(df.columns), index=0)
    topn       = c4.slider("Top-N", 3, 50, 10)

    period = _norm_period_value(period_raw)
    rule   = RULE_MAP[period]

    # ============================= 2) LẬP SERIES THEO SCHEMA =============================
    if not time_col:
        st.warning("Vui lòng chọn cột thời gian.")
        st.stop()

    s_time = pd.to_datetime(df[time_col], errors="coerce")

    if schema == "Amount + Type (2-type cols)":
        if not (amt_col and type_col and adj_col):
            st.warning("Vui lòng chọn Amount, Txn type và Adj type.")
            st.stop()

        amt = pd.to_numeric(df[amt_col], errors="coerce").fillna(0.0)
        txn = df[type_col].astype(str)
        adj = df[adj_col].astype(str)

        def _isin(s, vals): return s.isin(set(map(str, vals))) if vals else pd.Series(False, index=s.index)

        # Txn masks
        m_tin     = _isin(txn, val_tin)
        m_tout    = _isin(txn, val_tout)
        m_returns = _isin(txn, val_returns)

        # Adj masks (dùng để tính Discount%)
        m_adj_sales = _isin(adj, val_adj_sales)
        m_adj_disc  = _isin(adj, val_adj_disc)

        # Amount theo nhóm
        sales_s   = amt.where(m_adj_sales, 0.0)   # doanh thu (adj = sales)
        disc_s    = amt.where(m_adj_disc,  0.0)   # chiết khấu (thường âm)
        tin_s     = amt.where(m_tin,  0.0)
        tout_s    = amt.where(m_tout, 0.0)
        returns_s = amt.where(m_returns, 0.0)

    else:
        def _num(col): return pd.to_numeric(df[col], errors="coerce").fillna(0.0) if col else pd.Series(0.0, index=df.index)
        sales_s   = _num(sales_col)
        returns_s = _num(returns_col)
        disc_s    = _num(disc_col)
        tin_s     = _num(tin_col)
        tout_s    = _num(tout_col)

    # Net = Sales + Transfer(in/out) − |Returns| − |Discount|
    transfer_s   = tin_s.abs() + tout_s.abs()
    sales_pos    = sales_s.abs()        # dùng để so sánh % với discount
    discount_abs = disc_s.abs()
    returns_abs  = returns_s.abs()
    net_s        = sales_s + transfer_s - returns_abs - discount_abs

    # ============================= 3) KPI — 2 × 4, gọn & cân đối =============================
    orders_total = (df[order_col].nunique() if order_col else len(df))
    prod_total   = (df.loc[(sales_pos + transfer_s) > 0, prod_col].nunique() if prod_col else np.nan)

    sales_total_pos    = float(sales_pos.sum())
    transfer_total_pos = float(transfer_s.sum())
    pos_total          = sales_total_pos + transfer_total_pos
    pct_sales          = (sales_total_pos/pos_total*100.0) if pos_total>0 else np.nan
    pct_transfer       = (transfer_total_pos/pos_total*100.0) if pos_total>0 else np.nan
    net_total          = float(net_s.sum())

    # Discount% = Σ|Discount| / Σ|Sales| (dựa Adj type → Sales/Discount)
    m_idx = s_time.dt.to_period("M").dt.start_time
    q_idx = s_time.dt.to_period("Q").dt.start_time

    mon_df = pd.DataFrame({"m": m_idx, "sales": sales_pos, "disc": discount_abs})
    q_df   = pd.DataFrame({"q": q_idx, "sales": sales_pos, "disc": discount_abs})

    q_agg = q_df.groupby("q").sum()
    disc_pct_quarter = (q_agg["disc"].iloc[-1] / q_agg["sales"].iloc[-1]) if (not q_agg.empty and q_agg["sales"].iloc[-1] > 0) else np.nan

    mon = mon_df.groupby("m").sum()
    mon = mon[mon["sales"] > 0]
    if not mon.empty:
        mon["year"] = mon.index.year
        cnt = mon.groupby("year").size()
        full_years = cnt[cnt >= 12].index.tolist()
        if full_years:
            ly = max(full_years)
            mon_ly = mon[mon["year"] == ly]
            disc_pct_month_avg = (mon_ly["disc"]/mon_ly["sales"]).mean()     # TB % theo tháng trong năm đầy đủ
            disc_pct_year      =  mon_ly["disc"].sum() / mon_ly["sales"].sum()
        else:
            disc_pct_month_avg = np.nan
            disc_pct_year      = np.nan
    else:
        disc_pct_month_avg = np.nan
        disc_pct_year      = np.nan

    # Row 1
    k1, k2, k3, k4 = st.columns(4)
    k1.metric("Net Sales", f"{net_total:,.0f}")
    k2.metric("Orders", f"{orders_total:,.0f}")
    k3.metric("Total product", f"{prod_total:,.0f}" if not np.isnan(prod_total) else "—")
    k4.metric("%Sales (of Sales+Transfer)", f"{pct_sales:.1f}%" if not np.isnan(pct_sales) else "—")
    # Row 2
    k5, k6, k7, k8 = st.columns(4)
    k5.metric("%Transfer (of Sales+Transfer)", f"{pct_transfer:.1f}%" if not np.isnan(pct_transfer) else "—")
    k6.metric("Discount% (last quarter)", f"{(disc_pct_quarter*100):.1f}%" if not np.isnan(disc_pct_quarter) else "—")
    k7.metric("Discount% avg monthly (last full year)", f"{(disc_pct_month_avg*100):.1f}%" if not np.isnan(disc_pct_month_avg) else "—")
    k8.metric("Discount% (last full year)", f"{(disc_pct_year*100):.1f}%" if not np.isnan(disc_pct_year) else "—")

    # ============================= 4) XU HƯỚNG — BAR + LINE (%Δ) =============================
    idx_p  = s_time.dt.to_period(PERIOD_MAP[rule]).dt.start_time
    ser    = (pd.DataFrame({"p": idx_p, "v": net_s})
                .groupby("p")["v"].sum()
                .asfreq(rule)
                .fillna(0.0))
    base   = ser.shift(1) if compare == "Prev" else ser.shift(YOY_LAG[rule])
    growth = (ser - base) / base.replace(0, np.nan)

    fig = go.Figure()
    fig.add_bar(x=ser.index, y=ser.values, name="Net Sales",
                text=[f"{v:,.0f}" for v in ser.values],
                textposition="outside", hoverinfo="skip")
    fig.add_scatter(x=growth.index, y=growth.values*100, name="%Δ",
                    mode="lines+markers+text",
                    text=[f"{(v*100):.1f}%" if pd.notna(v) else "" for v in growth.values],
                    textposition="top center",
                    line=dict(color="#F2C811", width=3), marker=dict(size=6),
                    hoverinfo="skip", yaxis="y2")
    fig.update_layout(
        barmode="overlay",
        xaxis_title=period,
        yaxis=dict(title="Net Sales"),
        yaxis2=dict(title="%Δ", overlaying="y", side="right", showgrid=False),
        margin=dict(l=10, r=10, t=10, b=10),
        hovermode=False, showlegend=True, height=420
    )
    st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})
    st.caption("Bar = doanh số theo Period; Line = %Δ so với baseline (Prev/YoY)")

    # ============================= 5) ĐÓNG GÓP THEO NHÓM (Top-N & Pie, có Filter) =============================
    st.markdown("### 🧱 Đóng góp theo nhóm")
    cL, cR = st.columns(2)

    # Measure cho đóng góp
    measure = st.radio("Y (measure)", ["Net Sales","Sales only","Returns","Discount"], horizontal=True)
    def _y_series(name):
        return {
            "Net Sales": net_s,
            "Sales only": sales_s,
            "Returns": returns_s,
            "Discount": disc_s,
        }.get(name, net_s)
    base_y = pd.to_numeric(_y_series(measure), errors="coerce").fillna(0.0)

    # Filter values cho Dimension (dời xuống đây; bỏ Count unique)
    if dim_col and dim_col != "—":
        dim_vals = df[dim_col].astype(str).fillna("(NA)")
        g_all_for_filter = (pd.DataFrame({"dim": dim_vals, "val": base_y})
                            .groupby("dim", dropna=False)["val"]
                            .sum()
                            .sort_values(ascending=False))
        options = list(g_all_for_filter.index)
        picked  = st.multiselect("Filter values (áp dụng cho Top-N & Pie)", options=options, default=options)
        mask_dim = dim_vals.isin(picked) if picked else pd.Series(True, index=df.index)
    else:
        mask_dim = pd.Series(True, index=df.index)

    with cL:
        if not dim_col or dim_col == "—":
            st.info("Chọn Dimension để xem Top-N.")
        else:
            g = (pd.DataFrame({"dim": df[dim_col].astype(str), "val": base_y})
                   .loc[mask_dim]
                   .groupby("dim", dropna=False)["val"]
                   .sum()
                   .sort_values(ascending=False))
            g_top = g.head(topn)
            total_all = max(g.sum(), 1e-12)
            cum_share = (g_top.cumsum() / total_all) * 100.0

            fig_top = go.Figure()
            fig_top.add_bar(x=g_top.index, y=g_top.values, name=measure,
                            text=[f"{v:,.0f}" for v in g_top.values],
                            textposition="outside", hoverinfo="skip")
            fig_top.add_scatter(x=g_top.index, y=cum_share.values, name="Cumulative %",
                                mode="lines+markers+text",
                                text=[f"{v:.1f}%" for v in cum_share.values],
                                textposition="top center", yaxis="y2", line=dict(color="#A0A0A0"))
            fig_top.update_layout(
                xaxis_title=dim_col, yaxis_title=measure,
                yaxis2=dict(title="Cumulative %", overlaying="y", side="right", showgrid=False),
                margin=dict(l=10,r=10,t=10,b=10), hovermode=False, showlegend=True, height=460
            )
            st.plotly_chart(fig_top, use_container_width=True, config={"displayModeBar": False})

    with cR:
        if not dim_col or dim_col == "—":
            st.info("Chọn Dimension để xem tỉ trọng.")
        else:
            g_all = (pd.DataFrame({"dim": df[dim_col].astype(str), "val": base_y})
                       .loc[mask_dim]
                       .groupby("dim", dropna=False)["val"]
                       .sum()
                       .sort_values(ascending=False))

            total_pos = float(g_all.clip(lower=0).sum())
            if total_pos <= 0:
                st.info("Không đủ giá trị dương để vẽ pie.")
            else:
                share = (g_all.clip(lower=0) / total_pos)
                if len(share) > topn:
                    top_share = share.head(topn)
                    other = max(0.0, 1.0 - float(top_share.sum()))
                    labels = list(top_share.index) + (["Other"] if other > 1e-9 else [])
                    values = list((top_share*100).round(2).values) + ([round(other*100,2)] if other > 1e-9 else [])
                else:
                    labels = list(share.index)
                    values = list((share*100).round(2).values)

                fig_pie = go.Figure(go.Pie(
                    labels=labels, values=values, hole=0.35, sort=False, direction="clockwise",
                    text=[f"{lbl} {val:.1f}%" for lbl, val in zip(labels, values)],
                    textinfo="text", hoverinfo="skip"
                ))
                fig_pie.update_layout(margin=dict(l=10,r=10,t=10,b=10), showlegend=False, height=460)
                st.plotly_chart(fig_pie, use_container_width=True, config={"displayModeBar": False})

    # ============================= 6) PHÂN BỔ THEO VÙNG/KÊNH =============================
    st.markdown("### 🗺️ Phân bổ theo Vùng/Kênh")
    rc_y = net_s
    if region_col:
        if channel_col:
            topn_ch = st.slider("Top-N Channel (stacked)", 3, 20, 5)
            ch_sum = pd.DataFrame({"ch": df[channel_col].astype(str), "v": rc_y}).groupby("ch")["v"].sum().sort_values(ascending=False)
            keep_channels = set(ch_sum.head(topn_ch).index)
            ch = df[channel_col].astype(str).where(df[channel_col].astype(str).isin(keep_channels), other="Other")
            g = pd.DataFrame({"Region": df[region_col].astype(str), "Channel": ch, "v": rc_y}).groupby(["Region","Channel"])["v"].sum().reset_index()
            piv = g.pivot(index="Region", columns="Channel", values="v").fillna(0.0)
            piv = piv.loc[piv.sum(axis=1).sort_values().index]
            fig_rc = go.Figure()
            for col in piv.columns:
                fig_rc.add_bar(x=piv.index, y=piv[col].values, name=str(col), hoverinfo="skip")
            fig_rc.update_layout(barmode="stack", xaxis_title="Region", yaxis_title="Net Sales",
                                 margin=dict(l=10,r=10,t=10,b=10), hovermode=False, showlegend=True, height=460)
            st.plotly_chart(fig_rc, use_container_width=True, config={"displayModeBar": False})
        else:
            reg = pd.DataFrame({"Region": df[region_col].astype(str), "v": rc_y}).groupby("Region")["v"].sum().sort_values(ascending=True)
            fig_r = go.Figure()
            fig_r.add_bar(x=reg.values, y=reg.index, orientation="h",
                          text=[f"{v:,.0f}" for v in reg.values], textposition="outside",
                          hoverinfo="skip", name="Net Sales")
            fig_r.update_layout(xaxis_title="Net Sales", yaxis_title="Region",
                                margin=dict(l=10,r=10,t=10,b=10), hovermode=False, showlegend=False, height=440)
            st.plotly_chart(fig_r, use_container_width=True, config={"displayModeBar": False})
    else:
        st.info("Chọn **Region** để xem phân bổ theo vùng. (Có thể thêm Channel để xem stacked)")

    # ============================= 7) BẢNG TỔNG HỢP =============================
    st.markdown("### 🧾 Bảng tổng hợp")
    tbl_mode = st.radio("Góc nhìn bảng", ["Theo kỳ","Theo dimension"], horizontal=True)

    def _fmt(x):
        if pd.isna(x): return "—"
        if isinstance(x,(int,float,np.integer,np.floating)): return f"{x:,.0f}"
        return str(x)

    if tbl_mode == "Theo kỳ":
        gg = pd.DataFrame({"p": idx_p, "v": net_s})
        agg = gg.groupby("p")["v"].agg(count="count", sum="sum", mean="mean", median="median").reset_index()
        tot = agg["sum"].sum()
        agg["share"] = np.where(tot!=0, agg["sum"]/tot, np.nan)
        if rule=="MS":   agg["Nhóm"] = agg["p"].dt.to_period("M").astype(str)
        elif rule=="QS": agg["Nhóm"] = agg["p"].dt.to_period("Q").astype(str)
        else:            agg["Nhóm"] = agg["p"].dt.year.astype(str)
        tbl = agg[["Nhóm"]].copy()
        tbl["Số dòng"]    = agg["count"].astype(int)
        tbl["Tổng"]       = agg["sum"].map(_fmt)
        tbl["Trung bình"] = agg["mean"].map(_fmt)
        tbl["Trung vị"]   = agg["median"].map(_fmt)
        tbl["Tỷ trọng"]   = (agg["share"]*100).round(2).map(lambda v: f"{v:.2f}%" if pd.notna(v) else "—")
        _to_num = pd.to_numeric(tbl["Tổng"].str.replace(",",""), errors="coerce")
        tbl = tbl.iloc[_to_num.sort_values(ascending=False).index]
    else:
        if not dim_col or dim_col == "—":
            st.info("Chọn Dimension để tổng hợp.")
            st.stop()
        gg = pd.DataFrame({"dim": df[dim_col].astype(str), "v": net_s})
        agg = gg.groupby("dim", dropna=False)["v"].agg(count="count", sum="sum", mean="mean", median="median").reset_index().rename(columns={"dim":"Nhóm"})
        tot = agg["sum"].sum()
        agg["share"] = np.where(tot!=0, agg["sum"]/tot, np.nan)
        tbl = agg[["Nhóm"]].copy()
        tbl["Số dòng"]    = agg["count"].astype(int)
        tbl["Tổng"]       = agg["sum"].map(_fmt)
        tbl["Trung bình"] = agg["mean"].map(_fmt)
        tbl["Trung vị"]   = agg["median"].map(_fmt)
        tbl["Tỷ trọng"]   = (agg["share"]*100).round(2).map(lambda v: f"{v:.2f}%" if pd.notna(v) else "—")
        _to_num = pd.to_numeric(tbl["Tổng"].str.replace(",",""), errors="coerce")
        tbl = tbl.iloc[_to_num.sort_values(ascending=False).index]

    st.dataframe(tbl, use_container_width=True, hide_index=True)

    # ---------------- 4) Dimension filter (tuỳ chọn) ----------------
    st.markdown("#### 🎛️ Lọc Dimension (X) cho biểu đồ (Optional) ")
    if dim_col and dim_col != "—":
        count_by_default = order_col or cust_col or prod_col
        count_by = st.selectbox("Count unique by", ["—"] + [c for c in [order_col, cust_col, prod_col] if c], index=0 if not count_by_default else 1)
        if count_by == "—": count_by = None

        dim_vals = df[dim_col].astype(str).fillna("(NA)")
        if count_by:
            cnt_series = df.groupby(dim_col)[count_by].nunique().sort_values(ascending=False)
        else:
            cnt_series = dim_vals.value_counts()

        labels = [f"{k} ({v:,})" for k, v in cnt_series.items()]
        picked = st.multiselect("Filter values", options=labels, default=labels)
        keep_values = set(x.rsplit(" (", 1)[0] for x in picked) if picked else set(cnt_series.index)
        m_dim = dim_vals.isin(keep_values)
    else:
        m_dim = pd.Series(True, index=df.index)

    # ---------------- 5) Trend: Bar + Line (%Δ) ----------------
    measure = st.radio("Y (measure)", ["Net Sales","Sales only","Returns","Discount"], horizontal=True)

    def _select_y_series(measure):
        mapping = {
            "Net Sales":   net_s,
            "Sales only":  sales_s,
            "Returns":     returns_s,
            "Discount":    disc_s,
        }
        return pd.to_numeric(mapping.get(str(measure), net_s), errors="coerce").fillna(0.0)

    base_y = _select_y_series(measure)
    idx_p = s_time.dt.to_period(PERIOD_MAP[RULE_MAP[_norm_period_value(period)]]).dt.start_time
    ser = (pd.DataFrame({"p": idx_p, "v": base_y}).groupby("p")["v"].sum().asfreq(RULE_MAP[period]).fillna(0.0))
    base = ser.shift(1) if compare == "Prev" else ser.shift(YOY_LAG[RULE_MAP[period]])
    growth = (ser - base) / base.replace(0, np.nan)

    fig = go.Figure()
    fig.add_bar(x=ser.index, y=ser.values, name=measure,
                text=[f"{v:,.0f}" for v in ser.values], textposition="outside", hoverinfo="skip")
    fig.add_scatter(x=growth.index, y=growth.values*100, name="%Δ",
                    mode="lines+markers+text",
                    text=[f"{(v*100):.1f}%" if pd.notna(v) else "" for v in growth.values],
                    textposition="top center",
                    line=dict(color="#F2C811", width=3), marker=dict(size=6),
                    hoverinfo="skip", yaxis="y2")
    fig.update_layout(barmode="overlay",
                      xaxis_title=period, yaxis=dict(title=measure),
                      yaxis2=dict(title="%Δ", overlaying="y", side="right", showgrid=False),
                      margin=dict(l=10, r=10, t=10, b=10),
                      showlegend=True, hovermode=False)
    st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})
    st.caption("Bar = doanh số theo Period; Line = %Δ so với baseline (Prev/YoY)")

# ---------------- 6) Đóng góp theo nhóm ----------------
    st.markdown("### 🧱 Đóng góp theo nhóm")
    cL, cR = st.columns(2)

    # Measure đang chọn cho đóng góp
    measure = st.radio("Y (measure)", ["Net Sales","Sales only","Returns","Discount"], horizontal=True)
    def _y_series(name):
        return {
            "Net Sales": net_s,
            "Sales only": sales_s,
            "Returns": returns_s,
            "Discount": disc_s,
        }.get(name, net_s)
    base_y = pd.to_numeric(_y_series(measure), errors="coerce").fillna(0.0)
    
    # --- Filter values (di chuyển xuống đây) ---
    if dim_col and dim_col != "—":
        dim_vals = df[dim_col].astype(str).fillna("(NA)")
        g_all = pd.DataFrame({"dim": dim_vals, "val": base_y}).groupby("dim", dropna=False)["val"].sum().sort_values(ascending=False)
        # danh sách giá trị để lọc (không “count unique by” nữa)
        options = list(g_all.index)
        default_sel = options  # mặc định chọn tất cả
        picked = st.multiselect("Filter values (áp dụng cho Top-N & Pie)", options=options, default=default_sel)
        mask_dim = dim_vals.isin(picked) if picked else pd.Series(True, index=df.index)
    else:
        mask_dim = pd.Series(True, index=df.index)
    
    with cL:
        if not dim_col or dim_col == "—":
            st.info("Chọn Dimension để xem Top-N.")
        else:
            g = (pd.DataFrame({"dim": df[dim_col].astype(str), "val": base_y})
                   .loc[mask_dim]
                   .groupby("dim", dropna=False)["val"]
                   .sum()
                   .sort_values(ascending=False))
            g_top = g.head(topn)
            total_all = max(g.sum(), 1e-12)
            cum_share = (g_top.cumsum() / total_all) * 100.0
    
            fig_top = go.Figure()
            fig_top.add_bar(x=g_top.index, y=g_top.values, name=measure,
                            text=[f"{v:,.0f}" for v in g_top.values], textposition="outside", hoverinfo="skip")
            fig_top.add_scatter(x=g_top.index, y=cum_share.values, name="Cumulative %",
                                mode="lines+markers+text",
                                text=[f"{v:.1f}%" for v in cum_share.values],
                                textposition="top center", yaxis="y2", line=dict(color="#A0A0A0"))
            fig_top.update_layout(xaxis_title=dim_col, yaxis_title=measure,
                                  yaxis2=dict(title="Cumulative %", overlaying="y", side="right", showgrid=False),
                                  margin=dict(l=10,r=10,t=10,b=10), hovermode=False, showlegend=True)
            st.plotly_chart(fig_top, use_container_width=True, config={"displayModeBar": False})
    
    with cR:
        if not dim_col or dim_col == "—":
            st.info("Chọn Dimension để xem tỉ trọng.")
        else:
            g_all = (pd.DataFrame({"dim": df[dim_col].astype(str), "val": base_y})
                       .loc[mask_dim]
                       .groupby("dim", dropna=False)["val"]
                       .sum()
                       .sort_values(ascending=False))
            total_pos = float(g_all.clip(lower=0).sum())
            if total_pos <= 0:
                st.info("Không đủ giá trị dương để vẽ pie.")
            else:
                share = (g_all.clip(lower=0) / total_pos)
                if len(share) > topn:
                    top_share = share.head(topn)
                    other = max(0.0, 1.0 - float(top_share.sum()))
                    labels = list(top_share.index) + (["Other"] if other > 1e-9 else [])
                    values = list((top_share*100).round(2).values) + ([round(other*100,2)] if other > 1e-9 else [])
                else:
                    labels = list(share.index)
                    values = list((share*100).round(2).values)
    
                fig_pie = go.Figure(go.Pie(
                    labels=labels, values=values, hole=0.35, sort=False, direction="clockwise",
                    text=[f"{lbl} {val:.1f}%" for lbl, val in zip(labels, values)],
                    textinfo="text", hoverinfo="skip"
                ))
                fig_pie.update_layout(margin=dict(l=10,r=10,t=10,b=10), showlegend=False)
                st.plotly_chart(fig_pie, use_container_width=True, config={"displayModeBar": False})

    # ---------------- 7) Phân bổ theo Vùng/Kênh (Net gồm Transfer) ----------------
    st.markdown("### 🗺️ Phân bổ theo Vùng/Kênh")
    rc_y = net_s  # Net đã gồm transfer
    if region_col:
        if channel_col:
            topn_ch = st.slider("Top-N Channel (stacked)", 3, 20, 5)
            ch_sum = pd.DataFrame({"ch": df[channel_col].astype(str), "v": rc_y}).groupby("ch")["v"].sum().sort_values(ascending=False)
            keep_channels = set(ch_sum.head(topn_ch).index)
            ch = df[channel_col].astype(str).where(df[channel_col].astype(str).isin(keep_channels), other="Other")
            g = pd.DataFrame({"Region": df[region_col].astype(str), "Channel": ch, "v": rc_y}).groupby(["Region","Channel"])["v"].sum().reset_index()
            piv = g.pivot(index="Region", columns="Channel", values="v").fillna(0.0)
            piv = piv.loc[piv.sum(axis=1).sort_values().index]
            fig_rc = go.Figure()
            for col in piv.columns:
                fig_rc.add_bar(x=piv.index, y=piv[col].values, name=str(col), hoverinfo="skip")
            fig_rc.update_layout(barmode="stack", xaxis_title="Region", yaxis_title="Net Sales",
                                 margin=dict(l=10,r=10,t=10,b=10), hovermode=False, showlegend=True, height=460)
            st.plotly_chart(fig_rc, use_container_width=True, config={"displayModeBar": False})
        else:
            reg = pd.DataFrame({"Region": df[region_col].astype(str), "v": rc_y}).groupby("Region")["v"].sum().sort_values(ascending=True)
            fig_r = go.Figure()
            fig_r.add_bar(x=reg.values, y=reg.index, orientation="h",
                          text=[f"{v:,.0f}" for v in reg.values], textposition="outside",
                          hoverinfo="skip", name="Net Sales")
            fig_r.update_layout(xaxis_title="Net Sales", yaxis_title="Region",
                                margin=dict(l=10,r=10,t=10,b=10), hovermode=False, showlegend=False, height=440)
            st.plotly_chart(fig_r, use_container_width=True, config={"displayModeBar": False})
    else:
        st.info("Chọn **Region** để xem phân bổ theo vùng. (Có thể thêm Channel để xem stacked)")

    # ---------------- 8) Bảng tổng hợp ----------------
    st.markdown("### 🧾 Bảng tổng hợp")
    tbl_mode = st.radio("Góc nhìn bảng", ["Theo kỳ","Theo dimension"], horizontal=True)

    def _fmt(x):
        if pd.isna(x): return "—"
        if isinstance(x,(int,float,np.integer,np.floating)): return f"{x:,.0f}"
        return str(x)

    if tbl_mode == "Theo kỳ":
        idx_p = s_time.dt.to_period(PERIOD_MAP[rule]).dt.start_time
        gg = pd.DataFrame({"p": idx_p, "v": net_s})
        agg = gg.groupby("p")["v"].agg(count="count", sum="sum", mean="mean", median="median").reset_index()
        tot = agg["sum"].sum()
        agg["share"] = np.where(tot!=0, agg["sum"]/tot, np.nan)
        if rule=="MS":   agg["Nhóm"] = agg["p"].dt.to_period("M").astype(str)
        elif rule=="QS": agg["Nhóm"] = agg["p"].dt.to_period("Q").astype(str)
        else:            agg["Nhóm"] = agg["p"].dt.year.astype(str)
        tbl = agg[["Nhóm"]].copy()
        tbl["Số dòng"]    = agg["count"].astype(int)
        tbl["Tổng"]       = agg["sum"].map(_fmt)
        tbl["Trung bình"] = agg["mean"].map(_fmt)
        tbl["Trung vị"]   = agg["median"].map(_fmt)
        tbl["Tỷ trọng"]   = (agg["share"]*100).round(2).map(lambda v: f"{v:.2f}%" if pd.notna(v) else "—")
        _to_num = pd.to_numeric(tbl["Tổng"].str.replace(",",""), errors="coerce")
        tbl = tbl.iloc[_to_num.sort_values(ascending=False).index]
    else:
        if not dim_col or dim_col == "—":
            st.info("Chọn Dimension để tổng hợp.")
            st.stop()
        gg = pd.DataFrame({"dim": df[dim_col].astype(str), "v": net_s}).loc[m_dim]
        agg = gg.groupby("dim", dropna=False)["v"].agg(count="count", sum="sum", mean="mean", median="median").reset_index().rename(columns={"dim":"Nhóm"})
        tot = agg["sum"].sum()
        agg["share"] = np.where(tot!=0, agg["sum"]/tot, np.nan)
        tbl = agg[["Nhóm"]].copy()
        tbl["Số dòng"]    = agg["count"].astype(int)
        tbl["Tổng"]       = agg["sum"].map(_fmt)
        tbl["Trung bình"] = agg["mean"].map(_fmt)
        tbl["Trung vị"]   = agg["median"].map(_fmt)
        tbl["Tỷ trọng"]   = (agg["share"]*100).round(2).map(lambda v: f"{v:.2f}%" if pd.notna(v) else "—")
        _to_num = pd.to_numeric(tbl["Tổng"].str.replace(",",""), errors="coerce")
        tbl = tbl.iloc[_to_num.sort_values(ascending=False).index]

    st.dataframe(tbl, use_container_width=True, hide_index=True)

with TAB2:
    st.subheader('🧪 Distribution & Shape')
    df = DF_FULL
    # === Rule Engine sync helper for Tab 2 ===
    def _sync_rule_engine_from_tab2(field: str, kind: str, rules: list[tuple]):
        """
        field: tên cột
        kind: 'numeric' | 'categorical'
        rules: list of tuples (rule_name, score, severity, detail)
        """
        # Lưu riêng cho Tab 2 (nếu muốn debug)
        SS.setdefault('rule_engine_tab2', {})
        SS['rule_engine_tab2'][field] = {'kind': kind, 'rules': rules}
    
        # Tổng hợp về "All Test" ở Tab Risk (giữ nguyên cấu trúc flags_from_tabs đang dùng)
        SS.setdefault('flags_from_tabs', [])
        # xóa entries cũ của Tab 2 cùng column (tránh trùng khi UI rerun)
        SS['flags_from_tabs'] = [
            r for r in SS['flags_from_tabs']
            if not (r.get('tab') == 'Distribution & Shape' and r.get('column') == field)
        ]
        # thêm entries mới
        SS['flags_from_tabs'].extend([
            {
                'tab': 'Distribution & Shape',
                'rule': name,
                'column': field,
                'score': float(score),
                'severity': severity,
                'detail': detail,
            }
            for (name, score, severity, detail) in rules
        ])

    if df is None or len(df) == 0:
        st.info('Chưa có dữ liệu. Hãy **Load full data** trước khi dùng TAB 2.')
    else:
        # Let user pick any field
        num_cols = [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]
        cat_cols = [c for c in df.columns if (not pd.api.types.is_numeric_dtype(df[c])) and (not pd.api.types.is_datetime64_any_dtype(df[c]))]

        if not num_cols and not cat_cols:
            st.info('Không tìm thấy cột phù hợp để phân phối.'); 
        else:
            c1, c2 = st.columns([1,1])
            with c1:
                field = st.selectbox('Chọn cột', options=(num_cols+cat_cols), key='ds_field')
            with c2:
                view = st.radio('Kiểu xem', ['Auto','Numeric only','Categorical only'], horizontal=True, key='ds_view')

            if pd.api.types.is_numeric_dtype(df[field]) and view in ['Auto','Numeric only']:
                bins = st.slider('Bins', 10, 100, max(10, min(60, SS.get('bins', 50))), 5, key='ds_bins')
                fig = px.histogram(df, x=field, nbins=int(bins), title=f'Distribution — {field}', text_auto=True)
                st_plotly(fig)
                fig2 = px.box(df, y=field, points=False, title=f'Spread — {field}')
                st_plotly(fig2)
                s = pd.to_numeric(df[field], errors='coerce').dropna()
                if not s.empty:
                    q1, q3 = s.quantile(0.25), s.quantile(0.75); iqr = q3-q1
                    out_lo, out_hi = q1-1.5*iqr, q3+1.5*iqr
                    out_rate = ((s<out_lo)|(s>out_hi)).mean()
                    st.caption(f'n={len(s):,} • mean={s.mean():,.2f} • sd={s.std():,.2f} • median={s.median():,.2f} • IQR={iqr:,.2f} • outliers≈{out_rate*100:.1f}%')
                     # === Numeric table (Distribution & Shape) + Rule context ===
                    num_col = field
                    s_num = pd.to_numeric(df[num_col], errors='coerce').replace([np.inf,-np.inf], np.nan).dropna()
                    if not s_num.empty:
                        desc = s_num.describe(percentiles=[.01,.05,.25,.50,.75,.95,.99]).to_dict()
                        p95 = desc.get('95%', np.nan); p99 = desc.get('99%', np.nan)
                        zero_ratio = float((s_num==0).mean())
                        skew = float(s_num.skew()) if len(s_num)>2 else np.nan
                        kurt = float(s_num.kurt()) if len(s_num)>3 else np.nan
                        try:
                            from scipy import stats as spstats
                            p_norm = float(spstats.normaltest(s_num)[1]) if len(s_num) >= 8 else np.nan
                        except Exception:
                            p_norm = np.nan
                    
                        stat_df = pd.DataFrame([{
                            'column': num_col,
                            'count': int(desc.get('count',0)),
                            'n_missing': int(len(df[num_col]) - int(desc.get('count',0))),
                            'mean': desc.get('mean'), 'std': desc.get('std'),
                            'min': desc.get('min'),  'p1': desc.get('1%'),  'p5': desc.get('5%'),
                            'q1': desc.get('25%'),  'median': desc.get('50%'), 'q3': desc.get('75%'),
                            'p95': desc.get('95%'), 'p99': desc.get('99%'), 'max': desc.get('max'),
                            'skew': skew, 'kurtosis': kurt, 'zero_ratio': zero_ratio,
                            'tail>p95': float((s_num>p95).mean()) if not np.isnan(p95) else None,
                            'tail>p99': float((s_num>p99).mean()) if not np.isnan(p99) else None,
                            'normality_p': (round(p_norm,4) if not np.isnan(p_norm) else None),
                        }])
                        st_df(stat_df, use_container_width=True, height=220)
                        # === Rule Engine (auto) — NUMERIC ===
                        _rules_num = []
                        # các ngưỡng có thể chỉnh:
                        ZERO_RATIO_TH = 0.10
                        TAIL_P99_TH   = 0.01
                        SKEW_ABS_TH   = 1.0
                        KURT_EXCESS_TH= 4.0
                        P_NORMAL_TH   = 0.05
                        
                        # các biến bạn đã có ở trên: num_col, s_num, p95, p99, zero_ratio, skew, kurt, p_norm
                        tail_p99 = float((s_num > p99).mean()) if not np.isnan(p99) else 0.0
                        
                        if not np.isnan(zero_ratio) and zero_ratio >= ZERO_RATIO_TH:
                            _rules_num.append(('ZERO_RATIO_HIGH', min(1.0, zero_ratio/0.50), 'MED', f'Zero ratio {zero_ratio:.2%} ≥ {int(ZERO_RATIO_TH*100)}%'))
                        
                        if tail_p99 >= TAIL_P99_TH:
                            _rules_num.append(('HEAVY_TAIL_GT_P99', min(1.0, tail_p99/0.10), 'MED', f'Tail >p99 ≈ {tail_p99:.2%}'))
                        
                        if not np.isnan(skew) and abs(skew) >= SKEW_ABS_TH:
                            _rules_num.append(('SKEW_HIGH', min(1.0, abs(skew)/3.0), 'MED', f'|skew|={abs(skew):.2f}'))
                        
                        if not np.isnan(kurt) and kurt >= KURT_EXCESS_TH:
                            _rules_num.append(('KURTOSIS_HIGH', min(1.0, kurt/10.0), 'MED', f'excess kurtosis={kurt:.2f}'))
                        
                        if not np.isnan(p_norm) and p_norm < P_NORMAL_TH:
                            _rules_num.append(('NON_NORMAL', 0.8, 'MED', f'normality p={p_norm:.4f} < {P_NORMAL_TH}'))
                        
                        # sync về Rule Engine (All Test)
                        _sync_rule_engine_from_tab2(field=num_col, kind='numeric', rules=_rules_num)
                        
                        # Hiển thị rule insight tại chỗ (auto)
                        if _rules_num:
                            st.caption('Rule insights (auto • numeric)')
                            st_df(pd.DataFrame([{'rule': r[0], 'score': f'{r[1]:.2f}', 'severity': r[2], 'detail': r[3]} for r in _rules_num]),
                                  use_container_width=True, height=160)
                        else:
                            st.caption('Rule insights (auto • numeric): none')


            elif view in ['Auto','Categorical only']:
                topn = st.number_input('Top N', 3, 50, 20, 1, key='ds_topn')
                topc = df[field].astype(str).value_counts(dropna=True).head(int(topn)).reset_index()
                topc.columns = [field, 'count']; topc['%'] = (topc['count']/topc['count'].sum())
                fig = px.bar(topc, x=field, y='count', title=f'Top {int(topn)} — {field}', text_auto=True)
                st_plotly(fig)
                # --- NEW: Summary table (categorical) ---
                series_cat = df[field].astype(str)
                n = int(series_cat.notna().sum())
                u = int(series_cat.nunique(dropna=True))
                top_vc = series_cat.value_counts(dropna=True)
                if not top_vc.empty:
                    top_val = top_vc.index[0]
                    top_freq = int(top_vc.iloc[0])
                    top_pct = (top_freq / n * 100) if n else 0.0
                else:
                    top_val, top_freq, top_pct = '—', 0, 0.0
                
                tbl_cat = pd.DataFrame({
                    'stat': ['count', 'unique', 'mode', 'freq', '%'],
                    'value': [n, u, top_val, top_freq, f'{top_pct:.1f}%']
                }).set_index('stat')
                st_df(tbl_cat, use_container_width=True)
                # === Rule Engine (auto) — CATEGORICAL ===
                _rules_cat = []
                DOM_RATIO_TH = 0.60
                HI_CARD_TH   = 1000
                
                dom_ratio = (top_freq / n) if n else 0.0
                
                if dom_ratio >= DOM_RATIO_TH:
                    _rules_cat.append(('CATEGORY_DOMINANCE', min(1.0, (dom_ratio - DOM_RATIO_TH)/0.40 + 0.5), 'MED', f'Top chiếm {dom_ratio:.1%}'))
                
                if u >= HI_CARD_TH:
                    _rules_cat.append(('HIGH_CARDINALITY', 0.6, 'MED', f'unique={u:,} ≥ {HI_CARD_TH:,}'))
                
                # sync về Rule Engine (All Test)
                _sync_rule_engine_from_tab2(field=field, kind='categorical', rules=_rules_cat)
                
                # Hiển thị rule insight tại chỗ (auto)
                if _rules_cat:
                    st.caption('Rule insights (auto • categorical)')
                    st_df(pd.DataFrame([{'rule': r[0], 'score': f'{r[1]:.2f}', 'severity': r[2], 'detail': r[3]} for r in _rules_cat]),
                          use_container_width=True, height=140)
                else:
                    st.caption('Rule insights (auto • categorical): none')
                

# ---- TAB 3: Test Correlation (explicit tests, typed selects, robust) ----
with TAB3:
    import re, numpy as np, pandas as pd
    from scipy import stats
    import plotly.express as px
    import plotly.graph_objects as go

    st.subheader("🧪 Test Correlation")

    if SS.get('df') is None or len(SS['df']) == 0:
        st.info("Hãy nạp dữ liệu trước.")
        st.stop()

    df = SS['df']
    all_cols = list(df.columns)

    # ---------- Type detectors ----------
    def tc_is_num(c):
        try: return pd.api.types.is_numeric_dtype(df[c])
        except: return False

    def tc_is_dt(c):
        if c not in df.columns: return False
        if pd.api.types.is_datetime64_any_dtype(df[c]): return True
        return bool(re.search(r'(date|time|ngày|thời gian)', str(c), flags=re.I))

    def tc_is_cat(c):
        return (not tc_is_num(c)) and (not tc_is_dt(c))

    def tc_type(col):
        return 'datetime' if tc_is_dt(col) else ('numeric' if tc_is_num(col) else 'categorical')

    NUM_COLS = [c for c in all_cols if tc_is_num(c)]
    CAT_COLS = [c for c in all_cols if tc_is_cat(c)]
    DT_COLS  = [c for c in all_cols if tc_is_dt(c)]

    # Labeled options with icons & unique count for categorical
    def badge(c):
        if tc_is_dt(c):      icon = "🗓"
        elif tc_is_num(c):   icon = "🔢"
        else:                icon = "🔤"
        hint = ""
        if tc_is_cat(c):
            try: hint = f" · {df[c].nunique(dropna=True)}u"
            except: pass
        return f"{icon} {c}{hint}"

    label_to_col = {badge(c): c for c in all_cols}
    NUM_LB = [badge(c) for c in NUM_COLS]
    CAT_LB = [badge(c) for c in CAT_COLS]
    DT_LB  = [badge(c) for c in DT_COLS]

    # ---------- Helpers ----------
    def tc_make_period(s: pd.Series, period_lbl: str):
        freq = {"Month":"MS","Quarter":"QS","Year":"YS"}.get(period_lbl, "MS")
        return pd.to_datetime(s, errors='coerce').dt.to_period({"MS":"M","QS":"Q","YS":"Y"}[freq]).dt.start_time

    def tc_topn_cat(s: pd.Series, n=10):
        vc = s.astype(str).fillna("NaN").value_counts()
        top = vc.index[:n].tolist()
        return s.astype(str).where(s.astype(str).isin(top), "Khác")

    def tc_corr_ratio(categories, values):
        """η (0..1); FIX: dùng DataFrame align, tránh IndexingError."""
        y = pd.to_numeric(values, errors='coerce')
        dfv = pd.DataFrame({'cat': pd.Categorical(categories), 'y': y}).dropna()
        if dfv['cat'].nunique(dropna=True) < 2 or len(dfv) < 3:
            return np.nan
        grp = dfv.groupby('cat', observed=True)['y']
        y_mean = float(dfv['y'].mean())
        ss_between = sum(g.size * (float(g.mean()) - y_mean) ** 2 for _, g in grp)
        ss_total   = float(((dfv['y'] - y_mean) ** 2).sum())
        if ss_total == 0: return 0.0
        eta2 = ss_between / ss_total
        return float(np.sqrt(max(0.0, eta2)))

    def tc_anova_p(categories, values):
        """ANOVA p; FIX: align index bằng DataFrame trước khi groupby."""
        y = pd.to_numeric(values, errors='coerce')
        dfv = pd.DataFrame({'cat': pd.Categorical(categories), 'y': y}).dropna()
        if dfv['cat'].nunique(dropna=True) < 2: return np.nan
        groups = [g.values for _, g in dfv.groupby('cat', observed=True)['y'] if len(g) > 1]
        if len(groups) < 2: return np.nan
        try:
            _, p = stats.f_oneway(*groups)
            return float(p)
        except Exception:
            return np.nan

    def tc_cramers_v(x, y):
        tab = pd.crosstab(x, y, dropna=False)
        if tab.values.sum() == 0 or min(tab.shape) < 2: return np.nan, np.nan, tab
        chi2, p, _, _ = stats.chi2_contingency(tab)
        n = tab.values.sum(); r, c = tab.shape
        phi2 = chi2 / n
        phi2corr = max(0, phi2 - ((c-1)*(r-1))/(n-1))
        rcorr = c - ((c-1)**2)/(n-1)
        ccorr = r - ((r-1)**2)/(n-1)
        V = np.sqrt(phi2corr / max(1e-12, min(rcorr-1, ccorr-1)))
        return float(V), float(p), tab

    def r_strength(abs_r):
        return "yếu" if abs_r < 0.3 else ("vừa" if abs_r < 0.5 else "mạnh")

    def eta_strength(eta):
        if np.isnan(eta): return "—"
        return "yếu" if eta < 0.10 else ("vừa" if eta < 0.24 else "mạnh")

    def V_strength(V):
        if np.isnan(V): return "—"
        return "yếu" if V < 0.3 else ("vừa" if V < 0.5 else "mạnh")

 # ---------- UI (compact, typed selectors) ----------
    cfg = st.container(border=True)
    with cfg:
        # Hàng 1: test + X + Y + Fast (+ Robust nếu NN)
        c0, c1, c2, c3, c4 = st.columns([1.25, 1.1, 1.1, 0.7, 0.7])
        test_choice = c0.selectbox(
            "Loại test",
            ["Numeric ↔ Numeric", "Numeric ↔ Categorical", "Categorical ↔ Categorical", "Trend (time series)"],
            index=0,
            help="Chọn rõ test để X/Y chỉ hiện cột phù hợp."
        )
        fast_mode = c3.toggle("⚡ Fast", value=(len(df) >= 200_000))
    
        # --- Selectors gọn: lọc theo test ---
        overlay_pts = 0; topn_cat = 10
        robust = False
        dt_col, period_lbl, trans, roll_w = None, "Month", "%Δ MoM", 6
    
        if test_choice == "Numeric ↔ Numeric":
            if len(NUM_LB) < 2:
                st.warning("Thiếu cột numeric.")
                st.stop()
            x_label = c1.selectbox("X", NUM_LB, key="tc_x_nn", label_visibility="visible")
            y_label = c2.selectbox("Y", [lb for lb in NUM_LB if lb != x_label], key="tc_y_nn", label_visibility="visible")
            x_col, y_col = label_to_col[x_label], label_to_col[y_label]
            robust = c4.toggle("Robust", value=False, help="Spearman cho dữ liệu lệch/outlier")
            # Tùy chọn mở rộng — để ngắn gọn mặc định
            with st.expander("⚙️ Tùy chọn", expanded=False):
                colA, colB = st.columns([1,1])
                overlay_pts = colA.slider("Overlay points", 0, 5000, 1200, step=300,
                                          help="Lấy mẫu điểm chấm đè lên heatmap.", key="tc_overlay")
    
        elif test_choice == "Numeric ↔ Categorical":
            if (not NUM_LB) or (not CAT_LB):
                st.warning("Cần ≥1 numeric và ≥1 categorical.")
                st.stop()
            x_label = c1.selectbox("Numeric", NUM_LB, key="tc_x_nc", label_visibility="visible")
            y_label = c2.selectbox("Categorical", CAT_LB, key="tc_y_nc", label_visibility="visible")
            num_col, cat_col = label_to_col[x_label], label_to_col[y_label]
            with st.expander("⚙️ Tùy chọn", expanded=False):
                topn_cat = st.slider("Top N category", 3, 30, 10, key="tc_topn")
    
        elif test_choice == "Categorical ↔ Categorical":
            if len(CAT_LB) < 2:
                st.warning("Thiếu cột categorical.")
                st.stop()
            x_label = c1.selectbox("X", CAT_LB, key="tc_x_cc", label_visibility="visible")
            y_label = c2.selectbox("Y", [lb for lb in CAT_LB if lb != x_label], key="tc_y_cc", label_visibility="visible")
            x_col, y_col = label_to_col[x_label], label_to_col[y_label]
            with st.expander("⚙️ Tùy chọn", expanded=False):
                topn_cat = st.slider("Top N category", 3, 30, 10, key="tc_topn_cc")
    
        else:  # Trend (time series)
            if len(NUM_LB) < 2 or not DT_LB:
                st.warning("Cần ≥2 numeric và ≥1 datetime.")
                st.stop()
            x_label = c1.selectbox("X", NUM_LB, key="tc_x_tr", label_visibility="visible")
            y_label = c2.selectbox("Y", [lb for lb in NUM_LB if lb != x_label], key="tc_y_tr", label_visibility="visible")
            dt_label = c4.selectbox("🗓", DT_LB, key="tc_dt_tr", label_visibility="collapsed",
                                    help="Cột thời gian")
            x_col, y_col, dt_col = label_to_col[x_label], label_to_col[y_label], label_to_col[dt_label]
            # Hàng 2 gọn cho tham số thời gian
            t1, t2, t3 = st.columns([1.0, 1.0, 1.0])
            period_lbl = t1.selectbox("Period", ["Month","Quarter","Year"], index=0, key="tc_period", label_visibility="visible")
            trans = t2.selectbox("Biến đổi", ["%Δ MoM","%Δ YoY","MA(3)","MA(6)"], index=0, key="tc_trans", label_visibility="visible")
            roll_w = t3.slider("Rolling r (W)", 3, 24, 6, key="tc_roll", label_visibility="visible")

    # ---------- ROUTING ----------
    if test_choice == "Numeric ↔ Numeric":
        x = pd.to_numeric(df[x_col], errors='coerce')
        y = pd.to_numeric(df[y_col], errors='coerce')
        m = x.notna() & y.notna()
        x, y = x[m], y[m]
        n = int(len(x))
        if n < 3:
            st.info("Không đủ dữ liệu.")
            st.stop()
        if robust:
            r_val, p_val = stats.spearmanr(x, y)
            r_name = "Spearman"
        else:
            r_val, p_val = stats.pearsonr(x, y)
            r_name = "Pearson"
        R2 = (r_val**2)
        # slope/intercept OLS (quick)
        try:
            slope, intercept = np.polyfit(x, y, 1)
        except Exception:
            slope, intercept = np.nan, np.nan

        st.markdown(f"**{r_name} r = {r_val:.3f}**  ·  p={p_val:.4g}  ·  n={n}  ·  R²={R2:.3f}  ·  slope≈{slope:.3g}")

        # Chart: density heatmap + optional sample overlay
        if fast_mode:
            fig = px.density_heatmap(pd.DataFrame({x_col:x, y_col:y}),
                                     x=x_col, y=y_col, nbinsx=60, nbinsy=60, histfunc="count")
            if overlay_pts > 0:
                samp = pd.DataFrame({x_col:x, y_col:y}).sample(min(overlay_pts, n), random_state=42)
                fig.add_trace(go.Scattergl(x=samp[x_col], y=samp[y_col], mode='markers',
                                           marker=dict(size=3), name="sample"))
        else:
            fig = px.scatter(pd.DataFrame({x_col:x, y_col:y}), x=x_col, y=y_col,
                             opacity=0.55, render_mode="webgl")
        st.plotly_chart(fig, use_container_width=True)
        st.success(f"💡 Kết luận: Tương quan {r_strength(abs(r_val))} ({'+' if r_val>=0 else '−'})")

        SS['last_corr'] = pd.DataFrame([[1.0, r_val],[r_val,1.0]], index=[x_col,y_col], columns=[x_col,y_col])

    elif test_choice == "Numeric ↔ Categorical":
        s_num = pd.to_numeric(df[num_col], errors='coerce')
        s_cat = tc_topn_cat(df[cat_col], n=topn_cat)
        eta = tc_corr_ratio(s_cat, s_num)
        p_val = tc_anova_p(s_cat, s_num)
        eta2 = (eta**2) if not np.isnan(eta) else np.nan
        st.markdown(f"**η = {eta:.3f}** (η²={eta2:.3f})  ·  ANOVA p={p_val:.4g}")

        # Aggregated bar (median ± IQR/2)
        g = pd.DataFrame({cat_col:s_cat, num_col:s_num}).dropna() \
                .groupby(cat_col)[num_col].agg(q1=lambda s: s.quantile(0.25),
                                               med='median', q3=lambda s: s.quantile(0.75)) \
                .reset_index().sort_values('med', ascending=False)
        g['err'] = (g['q3'] - g['q1']) / 2.0
        fig = go.Figure(go.Bar(x=g[cat_col], y=g['med'],
                               error_y=dict(array=g['err'], visible=True)))
        fig.update_layout(yaxis_title=f"{num_col} (median ± IQR/2)")
        st.plotly_chart(fig, use_container_width=True)

        top_grp = str(g.iloc[0][cat_col]) if len(g) else "—"
        st.success(f"💡 Kết luận: Ảnh hưởng {eta_strength(eta)}; nhóm cao nhất: **{top_grp}**")

        SS['last_corr'] = None

    elif test_choice == "Categorical ↔ Categorical":
        sX = tc_topn_cat(df[x_col], n=topn_cat).astype(str)
        sY = tc_topn_cat(df[y_col], n=topn_cat).astype(str)
        V, p, tab = tc_cramers_v(sX, sY)
        st.markdown(f"**Cramér’s V = {V:.3f}**  ·  χ² p={p:.4g}")

        perc = (tab / tab.values.sum()).astype(float)
        fig = px.imshow(perc, aspect='auto', labels=dict(x=y_col, y=x_col, color='Share'))
        st.plotly_chart(fig, use_container_width=True)

        # Top residual pairs (hỗ trợ quan điểm)
        try:
            expected = np.outer(perc.sum(axis=1), perc.sum(axis=0)) * perc.values.sum()
            resid = (tab.values - expected) / np.sqrt(expected + 1e-12)
            idxs = np.dstack(np.unravel_index(np.argsort(-np.abs(resid), axis=None), resid.shape))[0][:3]
            bullets = [f"- **{tab.index[i]} × {tab.columns[j]}** (resid≈{resid[i,j]:.2f})" for (i,j) in idxs]
            st.info("Cặp lệch nổi bật:\n" + "\n".join(bullets))
        except Exception:
            pass

        st.success(f"💡 Kết luận: Liên hệ {V_strength(V)}.")
        SS['last_corr'] = None

    else:  # Trend (time series)
        tmp = df[[dt_col, x_col, y_col]].copy()
        tmp[dt_col] = pd.to_datetime(tmp[dt_col], errors='coerce')
        tmp = tmp.dropna(subset=[dt_col])
        tmp['__PERIOD__'] = tc_make_period(tmp[dt_col], period_lbl)
        agg = tmp.groupby('__PERIOD__')[[x_col, y_col]].sum().sort_index()
        if agg.shape[0] < max(3, roll_w):
            st.info("Chưa đủ kỳ để tính rolling.")
            st.stop()

        if trans in ("%Δ MoM", "%Δ YoY"):
            kmap = {"Month": {"MoM":1, "YoY":12},
                    "Quarter":{"MoM":1, "YoY":4},
                    "Year":{"MoM":1, "YoY":1}}
            k = kmap[period_lbl]["YoY" if "YoY" in trans else "MoM"]
            tsX = agg[x_col].pct_change(k)
            tsY = agg[y_col].pct_change(k)
            lbl = trans
        else:
            w = int(re.findall(r"\d+", trans)[0])
            tsX = agg[x_col].rolling(w, min_periods=max(2, w//2)).mean()
            tsY = agg[y_col].rolling(w, min_periods=max(2, w//2)).mean()
            lbl = f"MA({w})"

        ser = pd.DataFrame({f"{x_col} ({lbl})": tsX, f"{y_col} ({lbl})": tsY}).dropna()
        nK = int(len(ser))
        r_val = float(ser.iloc[:,0].corr(ser.iloc[:,1], method='pearson'))
        st.markdown(f"**r = {r_val:.3f}**  ·  n_kỳ={nK}")

        roll_r = ser.iloc[:,0].rolling(roll_w).corr(ser.iloc[:,1])
        fig_r = px.line(roll_r.reset_index(), x="__PERIOD__", y=0,
                        labels={"__PERIOD__":"Kỳ", "0":"rolling r"})
        st.plotly_chart(fig_r, use_container_width=True)

        best_lag, best_abs = 0, -1
        for L in range(-6, 7):
            v = ser.iloc[:,0].corr(ser.iloc[:,1].shift(L))
            if pd.notna(v) and abs(v) > best_abs:
                best_abs, best_lag = abs(v), L
        st.success(f"💡 Kết luận: Đồng pha {r_strength(abs(r_val))}; lag tốt nhất **{best_lag}** (|r|={best_abs:.3f}).")

        SS['last_corr'] = pd.DataFrame([[1.0, r_val],[r_val,1.0]],
                                       index=[f"{x_col} {lbl}", f"{y_col} {lbl}"],
                                       columns=[f"{x_col} {lbl}", f"{y_col} {lbl}"])

# ------------------------------- TAB 3: Benford -------------------------------
with TAB4:
    for k in ['bf1_res','bf2_res','bf1_col','bf2_col']:
        if k not in SS: SS[k]=None
    st.subheader('🔢 Benford Law — 1D & 2D')
    base_df = DF_FULL
    if not NUM_COLS:
        st.info('Không có cột numeric để chạy Benford.')
    else:
        run_on_full = True
        data_for_benford = DF_FULL
        # info removed
        c1,c2 = st.columns(2)
        with c1:
            amt1 = st.selectbox('Amount (1D)', NUM_COLS, key='bf1_col')
            if st.button('Run Benford 1D', key='btn_bf1'):
                ok,msg = _benford_ready(data_for_benford[amt1])
                if not ok: st.warning(msg)
                else: SS['bf1_res']=_benford_1d(data_for_benford[amt1])
        with c2:
            default_idx = 1 if len(NUM_COLS)>1 else 0
            amt2 = st.selectbox('Amount (2D)', NUM_COLS, index=default_idx, key='bf2_col')
            if st.button('Run Benford 2D', key='btn_bf2'):
                ok,msg = _benford_ready(data_for_benford[amt2])
                if not ok: st.warning(msg)
                else: SS['bf2_res']=_benford_2d(data_for_benford[amt2])
        g1,g2 = st.columns(2)
        with g1:
            if SS.get('bf1_res'):
                r=SS['bf1_res']; tb, var, p, MAD = r['table'], r['variance'], r['p'], r['MAD']
                if HAS_PLOTLY:
                    fig1 = go.Figure(); fig1.add_trace(go.Bar(x=tb['digit'], y=tb['observed_p'], name='Observed'))
                    fig1.add_trace(go.Scatter(x=tb['digit'], y=tb['expected_p'], name='Expected', mode='lines', line=dict(color='#F6AE2D')))
                    src_tag = 'FULL' if (SS['df'] is not None and SS.get('bf_use_full')) else 'SAMPLE'
                    fig1.update_layout(title=f'Benford 1D — Obs vs Exp ({SS.get("bf1_col")}, {src_tag})', height=340)
                    st_plotly(fig1)
                # --- Data quality — cột 1D đã chọn ---
                    _raw1 = data_for_benford[SS.get('bf1_col')]
                    _num1 = pd.to_numeric(_raw1, errors='coerce')
                    _total1 = len(_raw1)
                    _none_like1 = _raw1.astype('string').str.strip().str.lower().isin(['none','null']).sum()
                    _n_nan1  = _num1.isna().sum()
                    _n_zero1 = (_num1 == 0).sum()
                    _n_pos1  = (_num1 > 0).sum()
                    _n_neg1  = (_num1 < 0).sum()
                    _used1   = (_num1 != 0).sum()            
                    _base_clean1 = max(_total1 - _n_nan1 - _n_zero1, 0)
                    
                    qdf1 = pd.DataFrame({
                        'type': ['Total rows','NaN (numeric)','None/Null (text)','Zero (==0)',
                                 'Positive (>0)','Negative (<0)','Used for Benford (≠0)'],
                        'count': [int(_total1), int(_n_nan1), int(_none_like1), int(_n_zero1),
                                  int(_n_pos1), int(_n_neg1), int(_used1)]
                    })
                    qdf1['% vs total'] = (qdf1['count'] / _total1 * 100.0).round(2) if _total1>0 else 0.0
                    qdf1['% vs non-missing&non-zero'] = (
                        (qdf1['count'] / _base_clean1 * 100.0).round(2) if _base_clean1>0 else 0.0
                    )
                    st.caption('📋 Data quality — cột 1D đã chọn')
                    st_df(qdf1, use_container_width=True, height=180)
                    # --- Bảng % 1D (expected% / observed%) & diff% = observed% - expected% ---
                    color_thr_pct = 5.0  # drill-down theo chuẩn 5%
                    
                    t1 = pd.DataFrame({
                        'digit': tb['digit'].astype(int),
                        'expected_%': tb['expected_p'] * 100.0,
                        'observed_%': tb['observed_p'] * 100.0,
                    })
                    t1['diff_%'] = t1['observed_%'] - t1['expected_%']
                    
                    def _hl_percent1(v):
                        try:
                            return 'color: #d32f2f' if abs(float(v)) >= color_thr_pct else ''
                        except Exception:
                            return ''
                    
                    sty1 = (
                        t1.round(2)
                          .style
                          .format({'expected_%': '{:.2f}%', 'observed_%': '{:.2f}%', 'diff_%': '{:.2f}%'})
                          .applymap(_hl_percent1, subset=['diff_%'])
                    )
                    st_df(sty1, use_container_width=True, height=220)
                    
                    # --- Drill-down 1D cho những digit lệch ≥5% (tính theo diff_% ở trên) ---
                    bad_digits_1d = t1.loc[t1['diff_%'].abs() >= color_thr_pct, 'digit'].astype(int).tolist()
                    if bad_digits_1d:
                        with st.expander('🔎 Drill-down 1D: các chữ số lệch (|diff%| ≥ 5%)', expanded=False):
                            mode1 = st.radio('Chế độ hiển thị', ['Ngắn gọn','Xổ hết'], index=0,
                                             horizontal=True, key='bf1_drill_mode')
                    
                            import re as _re_local
                            def _digits_str(x):
                                xs = ("%.15g" % float(x))
                                return _re_local.sub(r"[^0-9]", "", xs).lstrip("0")
                            def _first1(v):
                                ds = _digits_str(v)
                                return int(ds[0]) if len(ds) >= 1 else np.nan
                    
                            s1_num = pd.to_numeric(data_for_benford[SS['bf1_col']], errors='coerce') \
                                       .replace([np.inf, -np.inf], np.nan).dropna().abs()
                            d1 = s1_num.apply(_first1).dropna()
                    
                            for dg in bad_digits_1d:
                                idx = d1[d1 == dg].index
                                st.markdown(f'**Digit {dg}** — {len(idx):,} rows')
                                if len(idx) == 0:
                                    continue
                                if mode1 == 'Xổ hết':
                                    st_df(data_for_benford.loc[idx].head(2000), use_container_width=True, height=260)
                                else:
                                    st_df(data_for_benford.loc[idx, [SS.get("bf1_col")]].head(200),
                                          use_container_width=True, height=220)
                    
                    # --- Thông điệp trạng thái dùng ngưỡng slider (so sánh theo tỷ lệ, không phải % point) ---
                    thr = SS.get('risk_diff_threshold', 0.05)               # ví dụ 0.05 = 5%
                    maxdiff_pp = float(t1['diff_%'].abs().max())            # % point
                    maxdiff_ratio = maxdiff_pp / 100.0                      # đổi về tỷ lệ để so với thr
                    
                    msg = '🟢 Green'
                    if maxdiff_ratio >= 2*thr:
                        msg = '🚨 Red'
                    elif maxdiff_ratio >= thr:
                        msg = '🟡 Yellow'
                    
                    sev = '🟢 Green'
                    if (p < 0.01) or (MAD > 0.015): sev = '🚨 Red'
                    elif (p < 0.05) or (MAD > 0.012): sev = '🟡 Yellow'
                    
                    st.info(f"Diff% status: {msg} • p={p:.4f}, MAD={MAD:.4f} ⇒ Benford severity: {sev}")

                    
        with g2:
            if SS.get('bf2_res'):
                r2=SS['bf2_res']; tb2, var2, p2, MAD2 = r2['table'], r2['variance'], r2['p'], r2['MAD']
                if HAS_PLOTLY:
                    fig2 = go.Figure(); fig2.add_trace(go.Bar(x=tb2['digit'], y=tb2['observed_p'], name='Observed'))
                    fig2.add_trace(go.Scatter(x=tb2['digit'], y=tb2['expected_p'], name='Expected', mode='lines', line=dict(color='#F6AE2D')))
                    src_tag = 'FULL' if (SS['df'] is not None and SS.get('bf_use_full')) else 'SAMPLE'
                    fig2.update_layout(title=f'Benford 2D — Obs vs Exp ({SS.get("bf2_col")}, {src_tag})', height=340)
                    st_plotly(fig2)
                # --- Data quality — cột 2D đã chọn ---
                    _raw2 = data_for_benford[SS.get('bf2_col')]
                    _num2 = pd.to_numeric(_raw2, errors='coerce')
                    _total2 = len(_raw2)
                    _none_like2 = _raw2.astype('string').str.strip().str.lower().isin(['none','null']).sum()
                    _n_nan2  = _num2.isna().sum()
                    _n_zero2 = (_num2 == 0).sum()
                    _n_pos2  = (_num2 > 0).sum()
                    _n_neg2  = (_num2 < 0).sum()
                    _used2   = (_num2 != 0).sum()            # Used for Benford: > 0 (giữ đúng logic tab này)
                    _base_clean2 = max(_total2 - _n_nan2 - _n_zero2, 0)
                    
                    qdf2 = pd.DataFrame({
                        'type': ['Total rows','NaN (numeric)','None/Null (text)','Zero (==0)',
                                 'Positive (>0)','Negative (<0)','Used for Benford (≠0)'],
                        'count': [int(_total2), int(_n_nan2), int(_none_like2), int(_n_zero2),
                                  int(_n_pos2), int(_n_neg2), int(_used2)]
                    })
                    qdf2['% vs total'] = (qdf2['count'] / _total2 * 100.0).round(2) if _total2>0 else 0.0
                    qdf2['% vs non-missing&non-zero'] = (
                        (qdf2['count'] / _base_clean2 * 100.0).round(2) if _base_clean2>0 else 0.0
                    )
                    st.caption('📋 Data quality — cột 2D đã chọn')
                    st_df(qdf2, use_container_width=True, height=180)
                    
                   # --- Bảng % 2D (expected% / observed%) & diff% = observed% - expected% ---
                    color_thr_pct = 5.0  # drill-down theo chuẩn 5%
                    
                    t2 = pd.DataFrame({
                        'digit': tb2['digit'].astype(int),
                        'expected_%': tb2['expected_p'] * 100.0,
                        'observed_%': tb2['observed_p'] * 100.0,
                    })
                    t2['diff_%'] = t2['observed_%'] - t2['expected_%']
                    
                    def _hl_percent2(v):
                        try:
                            return 'color: #d32f2f' if abs(float(v)) >= color_thr_pct else ''
                        except Exception:
                            return ''
                    
                    sty2 = (
                        t2.round(2)
                          .style
                          .format({'expected_%': '{:.2f}%', 'observed_%': '{:.2f}%', 'diff_%': '{:.2f}%'})
                          .applymap(_hl_percent2, subset=['diff_%'])
                    )
                    st_df(sty2, use_container_width=True, height=220)
                    
                    # --- Drill-down 2D cho những digit lệch ≥5% (tính theo diff_% ở trên) ---
                    bad_digits_2d = t2.loc[t2['diff_%'].abs() >= color_thr_pct, 'digit'].astype(int).tolist()
                    if bad_digits_2d:
                        with st.expander('🔎 Drill-down 2D: các chữ số lệch (|diff%| ≥ 5%)', expanded=False):
                            mode2 = st.radio('Chế độ hiển thị', ['Ngắn gọn','Xổ hết'], index=0,
                                             horizontal=True, key='bf2_drill_mode')
                    
                            import re as _re_local
                            def _digits_str(x):
                                xs = ("%.15g" % float(x))
                                return _re_local.sub(r"[^0-9]", "", xs).lstrip("0")
                            def _first2(v):
                                ds = _digits_str(v)
                                return int(ds[:2]) if len(ds) >= 2 else (int(ds) if len(ds) == 1 and ds != '0' else np.nan)
                    
                            s2_num = pd.to_numeric(data_for_benford[SS['bf2_col']], errors='coerce') \
                                       .replace([np.inf, -np.inf], np.nan).dropna().abs()
                            d2 = s2_num.apply(_first2).dropna()
                    
                            for dg in bad_digits_2d:
                                idx = d2[d2 == dg].index
                                st.markdown(f'**Digit {dg}** — {len(idx):,} rows')
                                if len(idx) == 0:
                                    continue
                                if mode2 == 'Xổ hết':
                                    st_df(data_for_benford.loc[idx].head(2000), use_container_width=True, height=260)
                                else:
                                    st_df(data_for_benford.loc[idx, [SS.get("bf2_col")]].head(200),
                                          use_container_width=True, height=220)
                    
                    # --- Thông điệp trạng thái dùng ngưỡng slider (so sánh theo tỷ lệ, không phải % point) ---
                    thr = SS.get('risk_diff_threshold', 0.05)
                    maxdiff2_pp = float(t2['diff_%'].abs().max())
                    maxdiff2_ratio = maxdiff2_pp / 100.0
                    
                    msg2 = '🟢 Green'
                    if maxdiff2_ratio >= 2*thr:
                        msg2 = '🚨 Red'
                    elif maxdiff2_ratio >= thr:
                        msg2 = '🟡 Yellow'
                    
                    sev2 = '🟢 Green'
                    if (p2 < 0.01) or (MAD2 > 0.015): sev2 = '🚨 Red'
                    elif (p2 < 0.05) or (MAD2 > 0.012): sev2 = '🟡 Yellow'
                    
                    st.info(f"Diff% status: {msg2} • p={p2:.4f}, MAD={MAD2:.4f} ⇒ Benford severity: {sev2}")

# ------------------------------ TAB ? : Statistics Test (ANOVA & Nonparametric, balanced UI) ------------------------------
with TAB5:
    import numpy as np, pandas as pd, re
    import plotly.express as px
    import plotly.graph_objects as go
    from scipy import stats
    import streamlit as st

    st.subheader("📊 Statistics Test — ANOVA & Nonparametric")

    # ===== Data guard =====
    DF = SS.get('df')
    if DF is None or len(DF) == 0:
        st.info("Hãy nạp dữ liệu trước.")
        st.stop()

    # ===== Type helpers =====
    def is_num(c):
        try: return pd.api.types.is_numeric_dtype(DF[c])
        except: return False
    def is_dt(c):
        if c not in DF.columns: return False
        if pd.api.types.is_datetime64_any_dtype(DF[c]): return True
        return bool(re.search(r'(date|time|ngày|thời gian)', str(c), flags=re.I))
    def is_cat(c):
        return (not is_num(c)) and (not is_dt(c))

    NUM_COLS = [c for c in DF.columns if is_num(c)]
    CAT_COLS = [c for c in DF.columns if is_cat(c)]

    # ===== Small utils =====
    def topn_cat(s: pd.Series, n=10):
        vc = s.astype(str).fillna("NaN").value_counts()
        keep = vc.index[:n].tolist()
        return s.astype(str).where(s.astype(str).isin(keep), "Khác")

    def group_summary(y, g):
        """Return summary per group: n, mean, std, median, se, ci95(≈1.96*se)."""
        d = pd.DataFrame({"y": y, "g": g}).dropna()
        if d.empty: 
            return pd.DataFrame(columns=["group","n","mean","std","median","se","ci95"])
        agg = d.groupby("g")["y"].agg(n="count", mean="mean", std="std", median="median")
        agg["se"] = agg["std"] / np.sqrt(agg["n"].clip(lower=1))
        agg["ci95"] = 1.96 * agg["se"]
        out = agg.reset_index().rename(columns={"g":"group"})
        # fillna to avoid plotly errors
        return out.replace([np.inf, -np.inf], np.nan).fillna(0.0)

    def holm_bonferroni(pvals, labels):
        """Holm-Bonferroni adjust (two-sided)."""
        p = np.asarray(pvals, dtype=float)
        m = len(p)
        order = np.argsort(p)  # ascending
        adj = np.empty(m, dtype=float)
        running_max = 0.0
        for r, idx in enumerate(order):
            adj_val = (m - r) * p[idx]
            running_max = max(running_max, adj_val)
            adj[idx] = min(1.0, running_max)
        return pd.DataFrame({"pair": labels, "p_raw": p, "p_adj_holm": adj}).sort_values("p_adj_holm")

    def one_way_anova_fast(y, g):
        """One-way ANOVA via group sums. Return F, p, df1, df2, eta2, omega2, leve_p."""
        d = pd.DataFrame({"y": pd.to_numeric(y, errors="coerce"), "g": g}).dropna()
        if d["g"].nunique() < 2 or len(d) < 3:
            return np.nan, np.nan, np.nan, np.nan, np.nan, np.nan, np.nan
        # Levene (center='median' bền vững)
        try:
            levene_p = stats.levene(*[grp["y"].values for _, grp in d.groupby("g")], center="median").pvalue
        except Exception:
            levene_p = np.nan

        grp = d.groupby("g")["y"].agg(n="count", mean="mean")
        ssq = d.assign(y2=d["y"]**2).groupby("g")["y2"].sum()
        grand_mean = float(d["y"].mean())
        ssb = float((grp["n"] * (grp["mean"] - grand_mean) ** 2).sum())
        ssw = float((ssq - grp["n"] * (grp["mean"] ** 2)).sum())
        sst = float(((d["y"] - grand_mean) ** 2).sum())
        k = int(grp.shape[0]); n = int(d.shape[0])
        df1 = k - 1; df2 = max(n - k, 1)
        msb = ssb / max(df1, 1); msw = ssw / max(df2, 1)
        F = (msb / msw) if msw > 0 else np.inf
        p = 1 - stats.f.cdf(F, df1, df2) if np.isfinite(F) else 0.0
        eta2 = (ssb / sst) if sst > 0 else np.nan
        omega2 = ((ssb - df1 * msw) / (sst + msw)) if (sst + msw) > 0 else np.nan
        return float(F), float(p), float(df1), float(df2), float(eta2), float(omega2), float(levene_p)

    def kruskal_eps2(H, k, n):
        """Epsilon-squared for Kruskal–Wallis."""
        return float((H - (k - 1)) / (n - k)) if (n - k) > 0 else np.nan

    # ===== UI helpers (balanced & hints) =====
    def _dtype_name(col):
        if col is None: return "—"
        try:
            if pd.api.types.is_datetime64_any_dtype(DF[col]): return "datetime"
            if pd.api.types.is_numeric_dtype(DF[col]): return "numeric"
            return "categorical"
        except Exception:
            return "unknown"

    def _type_hint(label, col, expect):
        actual = _dtype_name(col)
        ok = (actual == expect)
        icon = "✅" if ok else "⚠️"
        st.caption(f"{icon} {label}: `{col}` · {actual} (yêu cầu: {expect})")

    def _cheatsheet_note(expanded: bool = False):
        with st.expander("📝 Xác định nhanh theo mục tiêu & dữ liệu", expanded=expanded):
            # Thu nhỏ chữ & khoảng cách
            st.markdown(
                """
                <style>
                  .mini-note p, .mini-note li { margin-bottom: 0.15rem; }
                  .mini-note h5 { margin: 0.2rem 0 0.4rem 0; font-size: 1rem; }
                  .mini-note { font-size: 0.92rem; line-height: 1.25; }
                </style>
                """,
                unsafe_allow_html=True
            )
    
            # Tóm tắt 1 dòng theo mục tiêu & dữ liệu
            st.markdown(
                """
                <div class="mini-note">
                <strong>➤ So sánh trung bình/median (Y numeric)</strong><br>
                <em>Independent (between):</em> 2 nhóm → <strong>Welch t-test</strong> (mặc định) / <em>Mann–Whitney U</em>; ≥3 nhóm → <strong>One-way ANOVA</strong> (var≈) / <strong>Welch ANOVA</strong> (var≠) / <em>Kruskal–Wallis</em>; 2 yếu tố → <strong>Two-way ANOVA</strong> (kiểm <em>interaction</em>); có biến kiểm soát → <strong>ANCOVA</strong>.<br>
                <em>Repeated (within):</em> 2 điều kiện → <strong>Paired t-test</strong> / <em>Wilcoxon</em>; ≥3 điều kiện → <strong>RM-ANOVA</strong> / <em>Friedman</em>.<br>
                <em>Big data:</em> dùng <strong>Max rows (fit)</strong>, <strong>Top-N groups</strong>, <strong>Fast charts</strong> (heatmap/violin), <strong>sample overlay</strong>.
                </div>
                """,
                unsafe_allow_html=True
            )
    
            # Hai cột: ANOVA vs Non-Parametric
            c1, c2 = st.columns(2, gap="small")
            with c1:
                st.markdown('<div class="mini-note"><h5>ANOVA (Parametric)</h5>', unsafe_allow_html=True)
                st.markdown(
                    """
                    <div class="mini-note">
                    <em>Independent (between):</em><br>
                    • <strong>One-way ANOVA</strong>: Y numeric + factor categorical (≥2), giả định gần chuẩn & phương sai gần bằng. Var≠ → <strong>Welch ANOVA</strong>.<br>
                    • <strong>Two-way ANOVA</strong>: Factor A, B; đọc <em>interaction</em> A×B trước khi kết luận main effects.<br>
                    • <em>Post-hoc</em>: <strong>Tukey HSD</strong> (var≈) / <strong>Games-Howell</strong> (var≠).<br><br>
                    <em>Repeated (within):</em><br>
                    • <strong>RM-ANOVA</strong>: kiểm <em>sphericity</em>; vi phạm → hiệu chỉnh <strong>Greenhouse–Geisser</strong>. <em>Post-hoc</em> Bonferroni.<br>
                    • 2 điều kiện lặp → <strong>Paired t-test</strong>.
                    </div>
                    """,
                    unsafe_allow_html=True
                )
                st.markdown('</div>', unsafe_allow_html=True)
    
            with c2:
                st.markdown('<div class="mini-note"><h5>Non-Parametric</h5>', unsafe_allow_html=True)
                st.markdown(
                    """
                    <div class="mini-note">
                    <em>Independent (between):</em><br>
                    • 2 nhóm → <strong>Mann–Whitney U</strong> (effect size <em>r≈|Z|/√N</em>).<br>
                    • ≥3 nhóm → <strong>Kruskal–Wallis</strong> (effect size <em>ε²</em>); <em>Post-hoc</em> <strong>Dunn + Holm</strong>.<br><br>
                    <em>Repeated (within):</em><br>
                    • 2 điều kiện → <strong>Wilcoxon signed-rank</strong> (effect size <em>r</em>).<br>
                    • ≥3 điều kiện → <strong>Friedman</strong> (effect size <em>Kendall’s W</em>); <em>Post-hoc</em> cặp-đôi Wilcoxon + Holm.
                    </div>
                    """,
                    unsafe_allow_html=True
                )
                st.markdown('</div>', unsafe_allow_html=True)

    # ===== Tabs =====
    tab_a, tab_np = st.tabs(["ANOVA (Parametric)", "Nonparametric"])
# ====================== ANOVA (Parametric) — Unified UI like Nonparametric ======================
    with tab_a:
        mode_a = st.radio("Testing", ["Independent (between)", "Repeated (within)"], horizontal=True, key="anova_mode")
    
        # ---------- Independent (between) ----------
        if mode_a == "Independent (between)":
            if len(NUM_COLS) == 0 or len(CAT_COLS) == 0:
                st.info("Cần tối thiểu 1 cột numeric (Y) và 1 cột categorical (factor).")
            else:
                # Header balanced: Left (design), Right (note)
                box_top = st.container(border=True)
                with box_top:
                    L, R = st.columns(2)
                    with L:
                        st.markdown("### ANOVA — Independent (between)")
                        y_col  = st.selectbox("🎯 Dependent (numeric)", NUM_COLS, key="av_y")
                        a_col  = st.selectbox("🏷️ Factor A (categorical)", CAT_COLS, key="av_a")
                        use_two = st.toggle("➕ Two-way ANOVA (thêm Factor B)", value=False, key="av_two")
                        b_col = None
                        if use_two:
                            b_choices = [c for c in CAT_COLS if c != a_col]
                            b_col = st.selectbox("🏷️ Factor B (categorical)", b_choices, key="av_b")
                        # type hints
                        _type_hint("Dependent", y_col, "numeric")
                        _type_hint("Factor A", a_col, "categorical")
                        if use_two and b_col:
                            _type_hint("Factor B", b_col, "categorical")
                    with R:
                        _cheatsheet_note()
    
                # Controls balanced
                box_ctl = st.container(border=True)
                with box_ctl:
                    L, R = st.columns(2)
                    with L:
                        topN_A = int(st.number_input("Top N nhóm (Factor A)", 2, 50, 10, step=1, key="av_topn_a"))
                        if use_two:
                            topN_B = int(st.number_input("Top N nhóm (Factor B)", 2, 50, 8, step=1, key="av_topn_b"))
                        show_ci  = st.checkbox("Hiện 95% CI", value=True, key="av_ci")
                        posthoc  = (not use_two) and st.checkbox("Pairwise (Holm adjust)", value=True, key="av_posthoc")
                    with R:
                        max_fit  = int(st.number_input("Max rows (fit)", 5_000, 2_000_000, 300_000, step=5_000, key="av_max"))
                        fast     = st.toggle("⚡ Fast", value=(len(DF) >= 300_000), key="av_fast")
                        chart_sample = st.number_input("Chart sample overlay", 0, 200_000, 10_000, step=1_000, key="av_samp")
                        run = st.button("▶️ Run ANOVA", use_container_width=True, key="av_run")
    
                # Compute & report
                if run:
                    if not use_two:
                        # ----- One-way ANOVA (giữ công thức nhanh) -----
                        sub = DF[[y_col, a_col]].copy()
                        if len(sub) > max_fit:
                            sub = sub.sample(n=max_fit, random_state=42)
                        sub[a_col] = topn_cat(sub[a_col], n=topN_A)
    
                        y = pd.to_numeric(sub[y_col], errors="coerce")
                        g = sub[a_col].astype(str)
                        ok = (~y.isna()) & (~g.isna())
                        y, g = y[ok], g[ok]
    
                        summ = group_summary(y, g).sort_values("mean", ascending=False)
                        st.dataframe(summ, use_container_width=True, hide_index=True)
                        if summ.shape[0] < 2 or len(y) < 3:
                            st.warning("Không đủ nhóm/hàng để chạy ANOVA.")
                            st.stop()
    
                        F, p, df1, df2, eta2, omega2, lev_p = one_way_anova_fast(y, g)
    
                        # metric cards
                        m1, m2, m3, m4 = st.columns(4)
                        m1.metric("F", f"{F:.3f}")
                        m2.metric("p-value", f"{p:.4g}")
                        m3.metric("η²", f"{eta2:.3f}" if not np.isnan(eta2) else "—")
                        m4.metric("ω²", f"{omega2:.3f}" if not np.isnan(omega2) else "—")
                        st.caption(f"Levene (phương sai bằng nhau) p = {lev_p:.4g}")
    
                        # chart
                        if fast or not show_ci:
                            fig = px.bar(summ, x="group", y="mean", labels={"group": a_col, "mean": f"Mean {y_col}"})
                        else:
                            fig = go.Figure(go.Bar(x=summ["group"], y=summ["mean"],
                                                   error_y=dict(type="data", array=summ["ci95"], visible=True)))
                            fig.update_layout(yaxis_title=f"{y_col} (mean ± 95% CI)")
                        st.plotly_chart(fig, use_container_width=True)
    
                        # post-hoc (Welch t-test + Holm)
                        if posthoc and summ.shape[0] >= 2:
                            groups = summ["group"].tolist()
                            pvals, labs = [], []
                            for i in range(len(groups)):
                                gi = groups[i]; xi = y[g == gi].values
                                for j in range(i+1, len(groups)):
                                    gj = groups[j]; xj = y[g == gj].values
                                    if len(xi) >= 2 and len(xj) >= 2:
                                        tt = stats.ttest_ind(xi, xj, equal_var=False)
                                        pvals.append(float(tt.pvalue)); labs.append(f"{gi} vs {gj}")
                            if pvals:
                                adj = holm_bonferroni(np.array(pvals), np.array(labs))
                                diffs = []
                                for pair in adj["pair"]:
                                    gi, gj = str(pair).split(" vs ")
                                    mi = summ.loc[summ["group"] == gi, "mean"].values[0]
                                    mj = summ.loc[summ["group"] == gj, "mean"].values[0]
                                    diffs.append(mi - mj)
                                adj["mean_diff"] = diffs
                                st.dataframe(adj.head(50), use_container_width=True, hide_index=True)
                                st.caption("Pairwise Welch t-test (Holm-adjusted).")
    
                        strength = ("yếu" if (np.isnan(eta2) or eta2 < 0.06) else ("vừa" if eta2 < 0.14 else "mạnh"))
                        best = str(summ.iloc[0]["group"]) if len(summ) else "—"
                        st.success(f"**Kết luận:** Khác biệt giữa các nhóm {strength} (η²={eta2:.2f}). Nhóm cao nhất: **{best}**.")
    
                    else:
                        # ----- Two-way ANOVA (OLS + anova_lm) -----
                        try:
                            import statsmodels.api as sm
                            import statsmodels.formula.api as smf
                        except Exception:
                            st.error("Two-way ANOVA cần `statsmodels`. Hãy cài đặt gói này.")
                            st.stop()
    
                        sub = DF[[y_col, a_col, b_col]].dropna().copy()
                        if len(sub) > max_fit:
                            sub = sub.sample(n=max_fit, random_state=42)
                        sub[a_col] = topn_cat(sub[a_col], n=topN_A)
                        sub[b_col] = topn_cat(sub[b_col], n=topN_B)
    
                        # rename for formula
                        d = sub.rename(columns={y_col: "Y", a_col: "A", b_col: "B"})
                        d["Y"] = pd.to_numeric(d["Y"], errors="coerce")
                        d = d.dropna(subset=["Y"])
                        if d["A"].nunique() < 2 or d["B"].nunique() < 2:
                            st.warning("Cần ≥2 mức cho mỗi factor sau khi Top-N.")
                            st.stop()
    
                        model = smf.ols("Y ~ C(A) + C(B) + C(A):C(B)", data=d).fit()
                        an_tbl = sm.stats.anova_lm(model, typ=2)  # sum_sq, df, F, PR(>F)
                        st.dataframe(an_tbl, use_container_width=True)
    
                        # partial η² cho từng hiệu ứng
                        if "Residual" in an_tbl.index and "sum_sq" in an_tbl.columns:
                            ss_res = float(an_tbl.loc["Residual", "sum_sq"])
                            def peta(row): 
                                ss = float(row["sum_sq"])
                                return ss / (ss + ss_res) if (ss + ss_res) > 0 else np.nan
                            peta_vals = an_tbl.apply(peta, axis=1)
                            pe = peta_vals.to_dict()
                        else:
                            pe = {}
    
                        # cards: A, B, A×B
                        def card_val(name, col):
                            if name in an_tbl.index:
                                Fv = an_tbl.loc[name, "F"]; pv = an_tbl.loc[name, "PR(>F)"]
                                ev = pe.get(name, np.nan)
                                col.metric(name.replace("C(","").replace(")",""), f"F={Fv:.2f}", f"p={pv:.3g}")
                                if not np.isnan(ev): col.caption(f"partial η² ≈ {ev:.3f}")
                            else:
                                col.metric(name, "—", "—")
    
                        c1, c2, c3 = st.columns(3)
                        card_val("C(A)", c1); card_val("C(B)", c2); card_val("C(A):C(B)", c3)
    
                        # summary means (bar grouped)
                        grp = d.groupby(["A","B"])["Y"].agg(n="count", mean="mean").reset_index()
                        fig = px.bar(grp, x="A", y="mean", color="B", barmode="group",
                                     labels={"A": a_col, "B": b_col, "mean": f"Mean {y_col}"})
                        st.plotly_chart(fig, use_container_width=True)
    
                        # kết luận
                        pA = float(an_tbl.loc["C(A)", "PR(>F)"]) if "C(A)" in an_tbl.index else np.nan
                        pB = float(an_tbl.loc["C(B)", "PR(>F)"]) if "C(B)" in an_tbl.index else np.nan
                        pI = float(an_tbl.loc["C(A):C(B)", "PR(>F)"]) if "C(A):C(B)" in an_tbl.index else np.nan
                        msg = []
                        if not np.isnan(pI) and pI < 0.05:
                            msg.append("**có tương tác A×B** (p<0.05) — nên đọc theo từng lát cắt.")
                        if not np.isnan(pA) and pA < 0.05:
                            msg.append("Factor **A** có ý nghĩa.")
                        if not np.isnan(pB) and pB < 0.05:
                            msg.append("Factor **B** có ý nghĩa.")
                        if not msg: msg = ["Chưa thấy hiệu ứng có ý nghĩa (p≥0.05)."]
                        st.success(" ; ".join(msg))
    
        # ---------- Repeated (within) ----------
        else:
            cand_id = [c for c in DF.columns if is_cat(c)]
            cand_factor = [c for c in CAT_COLS]
            if len(NUM_COLS) == 0 or len(cand_id) == 0 or len(cand_factor) == 0:
                st.info("Cần: 1 numeric (Y), 1 ID (subject), 1 categorical (condition).")
            else:
                # Header balanced
                box_top_r = st.container(border=True)
                with box_top_r:
                    L, R = st.columns(2)
                    with L:
                        st.markdown("### ANOVA — Repeated (within)")
                        y_col = st.selectbox("🎯 Y (numeric)", NUM_COLS, key="av_rep_y")
                        id_col = st.selectbox("🧑‍🤝‍🧑 ID (subject)", cand_id, key="av_rep_id")
                        cond_col = st.selectbox("🏷️ Condition (within)", cand_factor, key="av_rep_cond")
                        _type_hint("Y", y_col, "numeric")
                        _type_hint("ID", id_col, "categorical")
                        _type_hint("Condition", cond_col, "categorical")
                    with R:
                        _cheatsheet_note()
    
                box_ctl_r = st.container(border=True)
                with box_ctl_r:
                    L, R = st.columns(2)
                    with L:
                        max_subj_fit = int(st.number_input("Max subjects (fit)", 50, 50_000, 5_000, step=50, key="av_rep_max"))
                    with R:
                        plot_subj = int(st.number_input("Spaghetti sample", 0, 1000, 80, step=20, key="av_rep_sp"))
                        run = st.button("▶️ Run", use_container_width=True, key="av_rep_run")
    
                if run:
                    # Prefer RM-ANOVA via statsmodels; nếu không có, hướng dẫn dùng Friedman (đã có ở tab Nonparametric)
                    try:
                        from statsmodels.stats.anova import AnovaRM
                    except Exception:
                        st.error("RM-ANOVA cần `statsmodels`. Bạn có thể dùng tab **Nonparametric → Friedman** như một thay thế.")
                        st.stop()
    
                    d0 = DF[[y_col, id_col, cond_col]].dropna().copy()
                    cnt = d0.groupby([id_col, cond_col]).size().unstack(cond_col).dropna()
                    keep_ids = cnt.index
                    d = d0[d0[id_col].isin(keep_ids)]
                    # limit subjects
                    uniq_ids = d[id_col].unique()
                    if len(uniq_ids) > max_subj_fit:
                        keep = pd.Index(uniq_ids).sample(max_subj_fit, random_state=42)
                        d = d[d[id_col].isin(keep)]
    
                    if d.empty or d[cond_col].nunique() < 2:
                        st.warning("Không đủ subject/điều kiện để chạy RM-ANOVA.")
                        st.stop()
    
                    model = AnovaRM(d, depvar=y_col, subject=id_col, within=[cond_col])
                    res = model.fit()
                    st.text(res.summary())
    
                    # Means + spaghetti
                    pivot = d.pivot_table(index=id_col, columns=cond_col, values=y_col, aggfunc="mean")
                    levels = list(pivot.columns)
                    means = pivot.mean().reset_index()
                    means.columns = ["cond","mean"]
                    fig = px.line(means, x="cond", y="mean", markers=True)
                    st.plotly_chart(fig, use_container_width=True)
                    if plot_subj > 0 and pivot.shape[0] > 0:
                        samp = pivot.sample(min(plot_subj, pivot.shape[0]), random_state=42)
                        for _, row in samp.iterrows():
                            fig.add_trace(go.Scatter(x=levels, y=row.values, mode="lines", opacity=0.25, showlegend=False))
                        st.plotly_chart(fig, use_container_width=True)
    
                    # Quick read
                    st.success("**Kết luận:** xem p-value của within-factor trong bảng; p<0.05 ⇒ có khác biệt giữa các điều kiện.")

    # ====================== NONPARAMETRIC ======================
    with tab_np:
        mode = st.radio("Testing", ["Independent (between)", "Repeated (within)"], horizontal=True, key="np_mode")

        # ---------- Independent (between) ----------
        if mode == "Independent (between)":
            if len(NUM_COLS) == 0 or len(CAT_COLS) == 0:
                st.info("Cần 1 numeric (Y) và 1 categorical (group).")
            else:
                # Balanced header
                box_top_np = st.container(border=True)
                with box_top_np:
                    L, R = st.columns(2)
                    with L:
                        st.markdown("### Nonparametric — Independent")
                        y_col = st.selectbox("🎯 Y (numeric)", NUM_COLS, key="np_y")
                        g_col = st.selectbox("🏷️ Group (categorical)", CAT_COLS, key="np_g")
                        _type_hint("Y", y_col, "numeric")
                        _type_hint("Group", g_col, "categorical")
                    with R:
                        _cheatsheet_note()

                box_ctl_np = st.container(border=True)
                with box_ctl_np:
                    L, R = st.columns(2)
                    with L:
                        topN = int(st.number_input("Top N groups", 2, 50, 10, step=1, key="np_topn"))
                        fast = st.toggle("⚡ Fast", value=(len(DF) >= 300_000), key="np_fast")
                    with R:
                        max_fit = int(st.number_input("Max rows (fit)", 5_000, 2_000_000, 300_000, step=5_000, key="np_max"))
                        run = st.button("▶️ Run", use_container_width=True, key="np_run")

                if run:
                    sub = DF[[y_col, g_col]].copy()
                    if len(sub) > max_fit:
                        sub = sub.sample(max_fit, random_state=42)
                    sub[g_col] = topn_cat(sub[g_col], n=topN)
                    y = pd.to_numeric(sub[y_col], errors="coerce")
                    g = sub[g_col].astype(str)
                    ok = (~y.isna()) & (~g.isna())
                    y, g = y[ok], g[ok]

                    summ = group_summary(y, g).sort_values("median", ascending=False)
                    st.dataframe(summ, use_container_width=True, hide_index=True)

                    groups = [y[g == lv].values for lv in summ["group"]]
                    k = len(groups)
                    n = int(sum(len(arr) for arr in groups))

                    if k == 2:
                        # Mann–Whitney U
                        ures = stats.mannwhitneyu(groups[0], groups[1], alternative="two-sided")
                        p = float(ures.pvalue); U = float(ures.statistic)
                        # z from p (two-sided)
                        z = float(stats.norm.isf(p / 2.0)) if p > 0 else np.inf
                        r_eff = z / np.sqrt(n) if n > 0 and np.isfinite(z) else np.nan
                        st.markdown(f"**Mann–Whitney U**: U = {U:.3f}, p = {p:.4g}, r ≈ {r_eff:.3f}")

                        fig = px.violin(pd.DataFrame({g_col: g, y_col: y}), x=g_col, y=y_col,
                                        box=True, points=False)
                        st.plotly_chart(fig, use_container_width=True)

                        hi = str(summ.iloc[0]['group']) if len(summ) else "—"
                        level = ("mạnh" if (not np.isnan(r_eff) and r_eff >= 0.5)
                                 else "vừa" if (not np.isnan(r_eff) and r_eff >= 0.3) else "yếu")
                        st.success(f"**Kết luận:** Khác biệt {level} (r≈{r_eff:.2f}). Nhóm median cao nhất: **{hi}**.")
                    else:
                        # Kruskal–Wallis
                        H, p = stats.kruskal(*groups)
                        eps2 = kruskal_eps2(H, k, n)
                        st.markdown(f"**Kruskal–Wallis**: H = {H:.3f}, p = {p:.4g}, ε² = {eps2:.3f}")

                        fig = go.Figure(go.Bar(x=summ["group"], y=summ["median"],
                                               error_y=dict(array=summ["ci95"], visible=True)))
                        fig.update_layout(yaxis_title=f"{y_col} (median ± 95% CI≈)")
                        st.plotly_chart(fig, use_container_width=True)

                        # Post-hoc: pairwise Mann–Whitney + Holm
                        pvals, labs = [], []
                        for i in range(k):
                            for j in range(i+1, k):
                                u = stats.mannwhitneyu(groups[i], groups[j], alternative="two-sided")
                                pvals.append(float(u.pvalue))
                                labs.append(f"{summ['group'].iloc[i]} vs {summ['group'].iloc[j]}")
                        if pvals:
                            adj = holm_bonferroni(np.array(pvals), np.array(labs))
                            st.dataframe(adj.head(50), use_container_width=True, hide_index=True)
                            st.caption("Pairwise Mann–Whitney (Holm-adjusted).")

                        strength = ("yếu" if (np.isnan(eps2) or eps2 < 0.06)
                                    else ("vừa" if eps2 < 0.14 else "mạnh"))
                        hi = str(summ.iloc[0]["group"]) if len(summ) else "—"
                        st.success(f"**Kết luận:** Khác biệt {strength} (ε²={eps2:.2f}). Nhóm median cao nhất: **{hi}**.")

        # ---------- Repeated (within) ----------
        else:
            cand_id = [c for c in DF.columns if is_cat(c)]
            cand_factor = [c for c in CAT_COLS]
            if len(NUM_COLS) == 0 or len(cand_id) == 0 or len(cand_factor) == 0:
                st.info("Cần: 1 numeric (Y), 1 ID (subject), 1 categorical (condition).")
            else:
                box_top_r = st.container(border=True)
                with box_top_r:
                    L, R = st.columns(2)
                    with L:
                        st.markdown("### Nonparametric — Repeated (within)")
                        y_col = st.selectbox("🎯 Y (numeric)", NUM_COLS, key="rep_y")
                        id_col = st.selectbox("🧑‍🤝‍🧑 ID (subject)", cand_id, key="rep_id")
                        cond_col = st.selectbox("🏷️ Condition (within)", cand_factor, key="rep_cond")
                        _type_hint("Y", y_col, "numeric")
                        _type_hint("ID", id_col, "categorical")
                        _type_hint("Condition", cond_col, "categorical")
                    with R:
                        _cheatsheet_note()

                box_ctl_r = st.container(border=True)
                with box_ctl_r:
                    L, R = st.columns(2)
                    with L:
                        max_subj_fit = int(st.number_input("Max subjects (fit)", 50, 50_000, 5_000, step=50, key="rep_max"))
                    with R:
                        plot_subj = int(st.number_input("Spaghetti sample", 0, 1000, 80, step=20, key="rep_sp"))
                        run = st.button("▶️ Run", use_container_width=True, key="rep_run")

                if run:
                    d0 = DF[[y_col, id_col, cond_col]].dropna().copy()
                    # Keep subjects having all levels
                    count = d0.groupby([id_col, cond_col]).size().unstack(cond_col).dropna()
                    subj_keep = count.index
                    d = d0[d0[id_col].isin(subj_keep)]

                    uniq_ids = d[id_col].unique()
                    if len(uniq_ids) > max_subj_fit:
                        keep = pd.Index(uniq_ids).sample(max_subj_fit, random_state=42)
                        d = d[d[id_col].isin(keep)]

                    pivot = d.pivot_table(index=id_col, columns=cond_col, values=y_col, aggfunc="mean")
                    pivot = pivot.dropna(axis=0)
                    levels = list(pivot.columns)
                    m = len(levels); n = pivot.shape[0]

                    if m == 2:
                        a = pivot[levels[0]].values
                        b = pivot[levels[1]].values
                        res = stats.wilcoxon(a, b, zero_method="wilcox", correction=False, alternative="two-sided", mode="auto")
                        p = float(res.pvalue); Wstat = float(res.statistic)
                        z = float(stats.norm.isf(p/2.0)) if p > 0 else np.inf
                        r_eff = z / np.sqrt(n) if n > 0 and np.isfinite(z) else np.nan
                        st.markdown(f"**Wilcoxon signed-rank**: W = {Wstat:.3f}, p = {p:.4g}, r ≈ {r_eff:.3f}")

                        means = pivot.mean().reset_index()
                        means.columns = ["cond","mean"]
                        fig = px.line(means, x="cond", y="mean", markers=True)
                        st.plotly_chart(fig, use_container_width=True)
                        if plot_subj > 0:
                            samp = pivot.sample(min(plot_subj, pivot.shape[0]), random_state=42)
                            for _, row in samp.iterrows():
                                fig.add_trace(go.Scatter(x=levels, y=row.values, mode="lines", opacity=0.25, showlegend=False))
                            st.plotly_chart(fig, use_container_width=True)

                        level = ("mạnh" if (not np.isnan(r_eff) and r_eff >= 0.5)
                                 else "vừa" if (not np.isnan(r_eff) and r_eff >= 0.3) else "yếu")
                        st.success(f"**Kết luận:** Khác biệt {level} (r≈{r_eff:.2f}).")

                    else:
                        fr = stats.friedmanchisquare(*[pivot[c].values for c in levels])
                        chi2 = float(fr.statistic); p = float(fr.pvalue)
                        W = chi2 / (n * m * (m + 1) / 12.0) if n > 0 else np.nan
                        st.markdown(f"**Friedman**: χ² = {chi2:.3f}, p = {p:.4g}, W = {W:.3f}")

                        means = pivot.mean().reset_index()
                        means.columns = ["cond","mean"]
                        fig = px.line(means, x="cond", y="mean", markers=True)
                        st.plotly_chart(fig, use_container_width=True)
                        if plot_subj > 0:
                            samp = pivot.sample(min(plot_subj, pivot.shape[0]), random_state=42)
                            for _, row in samp.iterrows():
                                fig.add_trace(go.Scatter(x=levels, y=row.values, mode="lines", opacity=0.25, showlegend=False))
                            st.plotly_chart(fig, use_container_width=True)

                        # Post-hoc: pairwise Wilcoxon + Holm
                        pvals, labs = [], []
                        for i in range(m):
                            for j in range(i+1, m):
                                wi = stats.wilcoxon(pivot[levels[i]], pivot[levels[j]],
                                                    zero_method="wilcox", correction=False,
                                                    alternative="two-sided", mode="auto")
                                pvals.append(float(wi.pvalue))
                                labs.append(f"{levels[i]} vs {levels[j]}")
                        if pvals:
                            adj = holm_bonferroni(np.array(pvals), np.array(labs))
                            st.dataframe(adj.head(50), use_container_width=True, hide_index=True)
                            st.caption("Pairwise Wilcoxon (Holm-adjusted).")

                        strength = ("yếu" if (np.isnan(W) or W < 0.1) else ("vừa" if W < 0.3 else "mạnh"))
                        best = str(means.sort_values("mean", ascending=False).iloc[0]["cond"])
                        st.success(f"**Kết luận:** Khác biệt {strength} (W={W:.2f}). Điều kiện cao nhất: **{best}**.")



# ------------------------------ TAB 6: Regression (Compact • Big-data friendly) ------------------------------
with TAB6:
    st.subheader('📘 Regression (Liner/Logistic')

    # ===== Safe imports =====
    try:
        import numpy as np, pandas as pd, re
        from sklearn.model_selection import train_test_split, KFold, cross_val_score
        from sklearn.linear_model import LinearRegression, Ridge, Lasso, LogisticRegression
        from sklearn.preprocessing import StandardScaler
        from sklearn.metrics import (
            r2_score, mean_squared_error, mean_absolute_error,
            accuracy_score, roc_auc_score, roc_curve
        )
        import plotly.express as px
        import plotly.graph_objects as go
    except Exception as e:
        st.error(f"Thiếu thư viện: {e}. Cài đặt scikit-learn / plotly trước khi dùng.")
        st.stop()

    DF = SS.get('df')
    if DF is None or len(DF) == 0:
        st.info("Hãy nạp dữ liệu trước.")
        st.stop()

    # ===== Type helpers =====
    def is_num(c):
        try: return pd.api.types.is_numeric_dtype(DF[c])
        except: return False
    def is_dt(c):
        if c not in DF.columns: return False
        if pd.api.types.is_datetime64_any_dtype(DF[c]): return True
        return bool(re.search(r'(date|time|ngày|thời gian)', str(c), flags=re.I))
    def is_cat(c):
        return (not is_num(c)) and (not is_dt(c))

    NUM_COLS = [c for c in DF.columns if is_num(c)]
    CAT_COLS = [c for c in DF.columns if is_cat(c)]

    # ===== Quick guide (collapsed) =====
    with st.expander("💡 Hướng dẫn chọn mô hình ", expanded=False):
        st.markdown(
            "- **Linear**: Target là **số liên tục** (Revenue, AOV…). Nếu lệch mạnh → bật **log1p(Y)**.\n"
            "- **Ridge/Lasso**: nhiều feature / đa cộng tuyến → ổn định hệ số.\n"
            "- **Logistic**: Target **nhị phân (0/1)** (Mua/Không, Fraud…). Mất cân bằng lớp → **class weight**.\n"
            "- **Big data**: dùng **Fast**, giới hạn **Max rows (fit)** và **Chart sample**."
        )

    tab_lin, tab_log = st.tabs(['Linear Regression', 'Logistic Regression'])

    # ============================== LINEAR ==============================
    with tab_lin:
        if len(NUM_COLS) < 2:
            st.info("Cần ≥2 biến numeric để chạy Linear.")
        else:
            # ---- Controls (compact) ----
            c1, c2, c3, c4, c5 = st.columns([1.2, 1.6, 0.8, 0.8, 0.8])
            y_lin = c1.selectbox("🎯 Target (numeric)", NUM_COLS, key="lin_y")
            X_cand = [c for c in NUM_COLS if c != y_lin]
            X_lin = c2.multiselect("🧩 Features (numeric)", options=X_cand,
                                   default=X_cand[:min(6, len(X_cand))], key="lin_X")
            test_size = c3.slider("Test %", 0.1, 0.5, 0.25, 0.05, key="lin_ts")
            fast_mode = c4.toggle("⚡ Fast", value=(len(DF) >= 300_000), key="lin_fast")
            run_lin = c5.button("▶️ Run", use_container_width=True, key="lin_run")

            with st.expander("⚙️ Advanced", expanded=False):
                a1, a2, a3, a4, a5 = st.columns(5)
                standardize = a1.checkbox("Standardize X", value=True, key="lin_std")
                logy = a1.checkbox("log1p(Y)", value=False, key="lin_logy")
                impute_na = a2.checkbox("Impute NA (median)", value=True, key="lin_impute")
                drop_const = a2.checkbox("Drop const", value=True, key="lin_const")
                penalty = a3.selectbox("Penalty", ["OLS", "Ridge", "Lasso"], index=0, key="lin_penalty")
                alpha = a3.slider("Alpha", 0.01, 10.0, 1.0, 0.01, key="lin_alpha")
                kcv = a4.slider("CV folds (train)", 3, 10, 5, key="lin_kcv")
                max_rows_fit = a5.number_input("Max rows (fit)", min_value=5_000, max_value=2_000_000,
                                               value=200_000, step=5_000, help="Giới hạn số dòng dùng để fit.")
                chart_sample = a5.number_input("Chart sample", min_value=0, max_value=200_000,
                                               value=10_000, step=1_000, help="0 = không overlay điểm")

            # ---- Fit & Report ----
            if run_lin:
                try:
                    sub = DF[[y_lin] + X_lin].copy()
                    # numeric coerce
                    for c in [y_lin] + X_lin:
                        if not pd.api.types.is_numeric_dtype(sub[c]):
                            sub[c] = pd.to_numeric(sub[c], errors="coerce")
                    # sample rows for fitting (speed)
                    if len(sub) > max_rows_fit:
                        sub = sub.sample(n=int(max_rows_fit), random_state=42)
                    # impute/drop
                    if impute_na:
                        med = sub[X_lin].median(numeric_only=True)
                        sub[X_lin] = sub[X_lin].fillna(med)
                        sub = sub.dropna(subset=[y_lin])
                    else:
                        sub = sub.dropna()
                    removed = []
                    if drop_const and len(X_lin) > 0:
                        nunique = sub[X_lin].nunique()
                        keep = [c for c in X_lin if nunique.get(c, 0) > 1]
                        removed = [c for c in X_lin if c not in keep]
                        X_lin = keep

                    if (len(sub) < (len(X_lin) + 5)) or (len(X_lin) == 0):
                        st.error("Không đủ dữ liệu sau xử lý (cần ≥ số features + 5).")
                    else:
                        X = sub[X_lin].copy()
                        y = sub[y_lin].copy()
                        y_t = np.log1p(y) if logy else y
                        Xtr, Xte, ytr, yte = train_test_split(X, y_t, test_size=test_size, random_state=42)

                        if standardize and Xtr.shape[1] > 0:
                            scaler = StandardScaler().fit(Xtr)
                            Xtr = pd.DataFrame(scaler.transform(Xtr), index=Xtr.index, columns=Xtr.columns)
                            Xte = pd.DataFrame(scaler.transform(Xte), index=Xte.index, columns=Xte.columns)

                        if penalty == "OLS":
                            model = LinearRegression()
                        elif penalty == "Ridge":
                            model = Ridge(alpha=alpha, random_state=42)
                        else:
                            model = Lasso(alpha=alpha, random_state=42, max_iter=10_000)

                        # CV (train) r2 — skip if fast_mode and too large features
                        cv_r2 = np.nan
                        if not fast_mode:
                            try:
                                cv = KFold(n_splits=kcv, shuffle=True, random_state=42)
                                cv_r2 = float(np.nanmean(cross_val_score(model, Xtr, ytr, cv=cv, scoring='r2')))
                            except Exception:
                                pass

                        model.fit(Xtr, ytr)
                        yhat_te = model.predict(Xte)
                        yhat = np.expm1(yhat_te) if logy else yhat_te
                        ytrue = np.expm1(yte) if logy else yte

                        r2 = r2_score(ytrue, yhat)
                        adj = 1 - (1 - r2) * (len(ytrue) - 1) / max(len(ytrue) - Xte.shape[1] - 1, 1)
                        rmse = float(np.sqrt(mean_squared_error(ytrue, yhat)))
                        mae = float(mean_absolute_error(ytrue, yhat))

                        # ===== Summary cards =====
                        m1, m2, m3, m4 = st.columns(4)
                        m1.metric("R² (test)", f"{r2:.3f}")
                        m2.metric("Adj R²", f"{adj:.3f}")
                        m3.metric("RMSE", f"{rmse:,.3f}")
                        m4.metric("MAE", f"{mae:,.3f}")
                        st.caption(f"CV R² (train): {cv_r2:.3f}" if not np.isnan(cv_r2) else "CV R² (train): —")

                        # ===== Equation / Coef =====
                        coef_s = pd.Series(model.coef_, index=X_lin, dtype=float)
                        intercept = float(model.intercept_)
                        with st.expander("📐 Phương trình hồi quy & hệ số", expanded=False):
                            st.code(
                                "Y{} = {:.6g} + ".format(" (log1p)" if logy else "", intercept) +
                                " + ".join([f"{b:.6g}·{name}" for name, b in coef_s.items()]),
                                language="text"
                            )
                            coef_show = coef_s.sort_values(key=lambda s: s.abs(), ascending=False).to_frame("coef")
                            st.dataframe(coef_show.head(30), use_container_width=True)

                        # ===== Charts in tabs (lightweight for big data) =====
                        t1, t2 = st.tabs(["Residuals & Fitted", "Feature importance"])
                        with t1:
                            # Residuals vs Fitted (density heatmap for fast) + hist residuals
                            resid = ytrue - yhat
                            if fast_mode:
                                df_plot = pd.DataFrame({"Fitted": yhat, "Residuals": resid})
                                fig1 = px.density_heatmap(df_plot, x="Fitted", y="Residuals", nbinsx=60, nbinsy=60)
                            else:
                                df_plot = pd.DataFrame({"Fitted": yhat, "Residuals": resid})
                                # overlay sample points if requested
                                if chart_sample > 0 and len(df_plot) > chart_sample:
                                    samp = df_plot.sample(chart_sample, random_state=42)
                                else:
                                    samp = df_plot
                                fig1 = px.scatter(samp, x="Fitted", y="Residuals", opacity=0.55, render_mode="webgl")
                            st.plotly_chart(fig1, use_container_width=True)

                            fig2 = px.histogram(resid, nbins=SS.get("bins", 50), title="Residuals distribution")
                            st.plotly_chart(fig2, use_container_width=True)

                        with t2:
                            imp = coef_s.abs().sort_values(ascending=False).head(20)
                            fig_imp = px.bar(x=imp.index, y=imp.values, labels={"x":"Feature","y":"|coef|"})
                            st.plotly_chart(fig_imp, use_container_width=True)

                        # ===== Conclusion =====
                        top_feat = imp.index[0] if len(imp) else "—"
                        strength = ("yếu" if r2 < 0.2 else ("vừa" if r2 < 0.5 else "mạnh"))
                        st.success(
                            f"**Kết luận:** Mô hình {penalty}{' (α='+str(alpha)+')' if penalty!='OLS' else ''} {strength} (R²={r2:.2f}). "
                            f"Yếu tố ảnh hưởng lớn nhất: **{top_feat}**. "
                            f"{'Đã log1p(Y). ' if logy else ''}{'Đã chuẩn hoá X. ' if standardize else ''}"
                            f"{'(Fast mode) ' if fast_mode else ''}"
                            f"{'Loại cột hằng: ' + ', '.join(removed[:5]) + ('…' if len(removed)>5 else '') if removed else ''}"
                        )
                except Exception as e:
                    st.error(f"Linear Regression error: {e}")

    # ============================== LOGISTIC ==============================
    with tab_log:
        # tìm các cột nhị phân
        bin_targets = []
        for c in DF.columns:
            s = DF[c].dropna()
            if s.nunique() == 2:
                bin_targets.append(c)

        if len(bin_targets) == 0:
            st.info("Không thấy cột nhị phân (2 giá trị).")
        else:
            # ---- Controls (compact) ----
            c1, c2, c3, c4, c5 = st.columns([1.2, 1.6, 0.8, 0.8, 0.8])
            y_col = c1.selectbox("🎯 Target (binary)", bin_targets, key="log_y")
            uniq = sorted(DF[y_col].dropna().unique().tolist())
            pos_label = c1.selectbox("Positive class", uniq, index=len(uniq)-1, key="log_pos")

            X_num_cand = [c for c in DF.columns if c != y_col and pd.api.types.is_numeric_dtype(DF[c])]
            X_cat_cand = [c for c in DF.columns if c != y_col and (not pd.api.types.is_numeric_dtype(DF[c])) and (not is_dt(c))]
            sel_num = c2.multiselect("🧩 Numeric features", options=X_num_cand, default=X_num_cand[:4], key="log_Xn")
            sel_cat = c2.multiselect("🏷️ Categorical features (optional)", options=X_cat_cand, default=[], key="log_Xc")

            test_size_l = c3.slider("Test %", 0.1, 0.5, 0.25, 0.05, key="log_ts")
            fast_mode_l = c4.toggle("⚡ Fast", value=(len(DF) >= 300_000), key="log_fast")
            run_log = c5.button("▶️ Run", use_container_width=True, key="log_run")

            with st.expander("⚙️ Advanced", expanded=False):
                a1, a2, a3, a4, a5 = st.columns(5)
                impute_na = a1.checkbox("Impute NA (median)", value=True, key="log_impute")
                drop_const = a1.checkbox("Drop const", value=True, key="log_const")
                class_bal = a2.checkbox("Class weight='balanced'", value=True, key="log_cw")
                standardize = a2.checkbox("Standardize numeric", value=True, key="log_std")
                topn_levels = a3.slider("Top-N / categorical", 3, 30, 8, key="log_topn")
                max_rows_fit = a4.number_input("Max rows (fit)", min_value=5_000, max_value=2_000_000,
                                               value=200_000, step=5_000, key="log_maxfit")
                chart_sample = a5.number_input("Chart sample", min_value=0, max_value=200_000,
                                               value=20_000, step=2_000, key="log_chartsamp")
                thr_mode = a3.selectbox("Gợi ý ngưỡng theo", ["F1","Youden J"], index=0, key="log_thrmode")
                thr_manual = a4.slider("Ngưỡng thủ công", 0.1, 0.9, 0.5, 0.05, key="log_thr")

            # ---- Fit & Report ----
            if run_log:
                try:
                    # y
                    y_raw = DF[y_col]
                    y = (y_raw == pos_label).astype(int)

                    # X numeric
                    Xn = DF[sel_num].copy()
                    for c in sel_num:
                        if not pd.api.types.is_numeric_dtype(Xn[c]):
                            Xn[c] = pd.to_numeric(Xn[c], errors="coerce")

                    # X categorical -> one-hot Top-N
                    Xc_list = []
                    for c in sel_cat:
                        s = DF[c].astype(str)
                        top = s.value_counts().head(topn_levels).index.tolist()
                        s2 = s.where(s.isin(top), "Other")
                        d = pd.get_dummies(s2, prefix=c, drop_first=True)
                        Xc_list.append(d)
                    Xc = pd.concat(Xc_list, axis=1) if Xc_list else pd.DataFrame(index=DF.index)

                    X_all = pd.concat([Xn, Xc], axis=1)

                    # limit rows for fit (speed)
                    idx = DF.index
                    if len(idx) > max_rows_fit:
                        idx = DF.sample(n=int(max_rows_fit), random_state=42).index
                    X_all = X_all.loc[idx]
                    y = y.loc[idx]

                    # impute / drop NA
                    if impute_na:
                        med = X_all.median(numeric_only=True)
                        X_all = X_all.fillna(med)
                    df_ready = pd.concat([y, X_all], axis=1).dropna()

                    # drop const
                    removed = []
                    if drop_const and X_all.shape[1] > 0:
                        nunique = df_ready.drop(columns=[y.name]).nunique()
                        keep = [c for c in X_all.columns if nunique.get(c, 0) > 1]
                        removed = [c for c in X_all.columns if c not in keep]
                        X_all = X_all[keep]
                        df_ready = pd.concat([y, X_all], axis=1)

                    if (len(df_ready) < (X_all.shape[1] + 10)) or (X_all.shape[1] == 0):
                        st.error("Không đủ dữ liệu sau xử lý (cần ≥ số features + 10).")
                    else:
                        X = df_ready.drop(columns=[y.name])
                        yb = df_ready[y.name]

                        Xtr, Xte, ytr, yte = train_test_split(X, yb, test_size=test_size_l, random_state=42, stratify=yb)

                        if standardize and Xtr.shape[1] > 0:
                            scaler = StandardScaler(with_mean=True, with_std=True).fit(Xtr)
                            Xtr = pd.DataFrame(scaler.transform(Xtr), index=Xtr.index, columns=Xtr.columns)
                            Xte = pd.DataFrame(scaler.transform(Xte), index=Xte.index, columns=Xte.columns)

                        model = LogisticRegression(max_iter=1000, class_weight=('balanced' if class_bal else None))
                        model.fit(Xtr, ytr)
                        proba = model.predict_proba(Xte)[:, 1]

                        # suggest threshold
                        thr_grid = np.linspace(0.1, 0.9, 17)
                        best_thr, best_score = 0.5, -1.0
                        for t in thr_grid:
                            pred = (proba >= t).astype(int)
                            tp = int(((pred == 1) & (yte == 1)).sum()); fp = int(((pred == 1) & (yte == 0)).sum())
                            fn = int(((pred == 0) & (yte == 1)).sum()); tn = int(((pred == 0) & (yte == 0)).sum())
                            prec = (tp / (tp + fp)) if (tp + fp) else 0.0
                            rec  = (tp / (tp + fn)) if (tp + fn) else 0.0
                            f1   = (2 * prec * rec / (prec + rec)) if (prec + rec) else 0.0
                            youden = (rec + (tn / (tn + fp) if (tn + fp) else 0.0) - 1.0)
                            score = f1 if (thr_mode == "F1") else youden
                            if score > best_score:
                                best_score, best_thr = score, float(t)

                        thr_use = float(thr_manual) if thr_manual else best_thr
                        pred = (proba >= thr_use).astype(int)

                        acc = accuracy_score(yte, pred)
                        try: auc = roc_auc_score(yte, proba)
                        except Exception: auc = np.nan

                        # ===== Summary cards =====
                        # Precision/Recall/F1 nhanh
                        tp = int(((pred == 1) & (yte == 1)).sum()); fp = int(((pred == 1) & (yte == 0)).sum())
                        fn = int(((pred == 0) & (yte == 1)).sum()); tn = int(((pred == 0) & (yte == 0)).sum())
                        prec = (tp / (tp + fp)) if (tp + fp) else 0.0
                        rec  = (tp / (tp + fn)) if (tp + fn) else 0.0
                        f1   = (2 * prec * rec / (prec + rec)) if (prec + rec) else 0.0

                        m1, m2, m3, m4 = st.columns(4)
                        m1.metric("Accuracy", f"{acc:.3f}")
                        m2.metric("Precision", f"{prec:.3f}")
                        m3.metric("Recall", f"{rec:.3f}")
                        m4.metric("F1", f"{f1:.3f}")
                        st.caption(f"AUC: {auc:.3f}" if not np.isnan(auc) else "AUC: —")
                        st.caption(f"Threshold dùng: {thr_use:.2f} (gợi ý: {best_thr:.2f} theo {thr_mode})")

                        # ===== Charts in tabs (lightweight) =====
                        t1, t2, t3 = st.tabs(["Confusion", "ROC", "Coefficients"])
                        with t1:
                            cm = pd.DataFrame([[tn, fp],[fn, tp]], index=["Actual 0","Actual 1"], columns=["Pred 0","Pred 1"])
                            fig_cm = px.imshow(cm, text_auto=True, aspect="auto", title="Confusion matrix")
                            st.plotly_chart(fig_cm, use_container_width=True)
                        with t2:
                            try:
                                fpr, tpr, _ = roc_curve(yte, proba)
                                # subsample ROC if huge
                                if fast_mode_l and len(fpr) > 50_000:
                                    step = len(fpr) // 50_000 + 1
                                    fpr, tpr = fpr[::step], tpr[::step]
                                fig_roc = px.area(x=fpr, y=tpr, labels={"x":"FPR","y":"TPR"}, title="ROC Curve")
                                fig_roc.add_shape(type="line", line=dict(dash="dash"), x0=0, x1=1, y0=0, y1=1)
                                st.plotly_chart(fig_roc, use_container_width=True)
                            except Exception:
                                st.info("Không vẽ được ROC.")
                        with t3:
                            try:
                                coef = pd.Series(model.coef_[0], index=Xtr.columns).sort_values(key=lambda s: s.abs(), ascending=False)
                                show = coef.head(20)
                                fig_coef = px.bar(x=show.index, y=show.values, labels={"x":"Feature","y":"coef"})
                                st.plotly_chart(fig_coef, use_container_width=True)
                                st.dataframe(show.head(30).to_frame("coef"), use_container_width=True)
                            except Exception:
                                st.info("Chưa có hệ số để hiển thị.")

                        # ===== Conclusion =====
                        strength = ("yếu" if (np.isnan(auc) or auc < 0.7) else ("vừa" if auc < 0.8 else "mạnh"))
                        try:
                            top_pos = coef[coef>0].index[0]
                            top_neg = coef[coef<0].index[0]
                            dir_feat = f" (+){top_pos}, (−){top_neg}"
                        except Exception:
                            dir_feat = ""
                        st.success(
                            f"**Kết luận:** Mô hình phân loại {strength} (F1={f1:.2f}, AUC={auc if not np.isnan(auc) else float('nan'):.2f})."
                            f"{' Tín hiệu mạnh:' + dir_feat if dir_feat else ''}  Ngưỡng: {thr_use:.2f}. "
                            f"{'(Fast mode) ' if fast_mode_l else ''}"
                            f"{'Đã chuẩn hoá numeric. ' if standardize else ''}"
                            f"{'Đã one-hot Top-N cho categorical. ' if len(sel_cat)>0 else ''}"
                            f"{'Loại cột hằng: ' + ', '.join(removed[:5]) + ('…' if len(removed)>5 else '') if 'removed' in locals() and removed else ''}"
                        )
                except Exception as e:
                    st.error(f"Logistic Regression error: {e}")
# -------------------------------- TAB 6: Flags --------------------------------
with TAB7:
    base_df = DF_FULL
    st.subheader('🚩 Fraud Flags')
    use_full_flags = True
    FLAG_DF = DF_FULL
    amount_col = st.selectbox('Amount (optional)', options=['(None)'] + NUM_COLS, key='ff_amt')
    dt_col = st.selectbox('Datetime (optional)', options=['(None)'] + DT_COLS, key='ff_dt')
    group_cols = st.multiselect('Composite key để dò trùng (tuỳ chọn)', options=[c for c in FLAG_DF.columns if (not SS.get('col_whitelist') or c in SS['col_whitelist'])], key='ff_groups')
    with st.expander('⚙️ Tham số quét cờ (điều chỉnh được)'):
        c1,c2,c3 = st.columns(3)
        with c1:
            thr_zero = st.number_input('Ngưỡng Zero ratio', 0.0, 1.0, 0.30, 0.05, key='ff_thr_zero')
            thr_tail99 = st.number_input('Ngưỡng Tail >P99 share', 0.0, 1.0, 0.02, 0.01, key='ff_thr_p99')
            thr_round = st.number_input('Ngưỡng .00/.50 share', 0.0, 1.0, 0.20, 0.05, key='ff_thr_round')
        with c2:
            thr_offh = st.number_input('Ngưỡng Off‑hours share', 0.0, 1.0, 0.15, 0.05, key='ff_thr_offh')
            thr_weekend = st.number_input('Ngưỡng Weekend share', 0.0, 1.0, 0.25, 0.05, key='ff_thr_weekend')
            dup_min = st.number_input('Số lần trùng key tối thiểu (≥)', 2, 100, 2, 1, key='ff_dup_min')
        with c3:
            near_str = st.text_input('Near approval thresholds (vd: 1,000,000; 2,000,000)', key='ff_near_list')
            near_eps_pct = st.number_input('Biên ±% quanh ngưỡng', 0.1, 10.0, 1.0, 0.1, key='ff_near_eps')
            use_daily_dups = st.checkbox('Dò trùng Amount theo ngày (khi có Datetime)', value=True, key='ff_dup_day')
        run_flags = st.button('🔎 Scan Flags', key='ff_scan', use_container_width=True)

    def _parse_near_thresholds(txt: str) -> list[float]:
        out=[]
        if not txt: return out
        for token in re.split(r"[;,]", txt):
            tok = token.strip().replace(',', '')
            if not tok: continue
            try: out.append(float(tok))
            except Exception: pass
        return out

    def _share_round_amounts(s: pd.Series) -> dict:
        x = pd.to_numeric(s, errors='coerce').dropna()
        if x.empty: return {'p_00': np.nan, 'p_50': np.nan}
        cents = (np.abs(x)*100).round().astype('Int64') % 100
        p00 = float((cents==0).mean()); p50 = float((cents==50).mean())
        return {'p_00': p00, 'p_50': p50}

    def _near_threshold_share(s: pd.Series, thresholds: list[float], eps_pct: float) -> pd.DataFrame:
        x = pd.to_numeric(s, errors='coerce').dropna()
        if x.empty or not thresholds: return pd.DataFrame(columns=['threshold','share'])
        eps = np.array(thresholds)*(eps_pct/100.0)
        res=[]
        for t,e in zip(thresholds, eps):
            if t<=0: continue
            share = float(((x >= (t-e)) & (x <= (t+e))).mean())
            res.append({'threshold': t, 'share': share})
        return pd.DataFrame(res)

    def compute_fraud_flags(df: pd.DataFrame, amount_col: Optional[str], datetime_col: Optional[str], group_id_cols: list[str], params: dict):
        flags=[]; visuals=[]
        num_cols2 = df.select_dtypes(include=[np.number]).columns.tolist()
        if num_cols2:
            zr_rows=[]
            for c in num_cols2:
                s = pd.to_numeric(df[c], errors='coerce')
                if len(s)==0: continue
                zero_ratio = float((s==0).mean()); zr_rows.append({'column':c, 'zero_ratio': round(zero_ratio,4)})
                if zero_ratio > params['thr_zero']:
                    flags.append({'flag':'High zero ratio','column':c,'threshold':params['thr_zero'],'value':round(zero_ratio,4),'note':'Threshold/rounding hoặc không sử dụng trường.'})
            if zr_rows: visuals.append(('Zero ratios (numeric)', pd.DataFrame(zr_rows).sort_values('zero_ratio', ascending=False)))
        amt = amount_col if (amount_col and amount_col!='(None)' and amount_col in df.columns) else None
        if amt:
            s_amt = pd.to_numeric(df[amt], errors='coerce').dropna()
            if len(s_amt)>20:
                p95=s_amt.quantile(0.95); p99=s_amt.quantile(0.99); tail99=float((s_amt>p99).mean())
                if tail99 > params['thr_tail99']:
                    flags.append({'flag':'Too‑heavy right tail (>P99)','column':amt,'threshold':params['thr_tail99'],'value':round(tail99,4),'note':'Kiểm tra outliers/segmentation/cut‑off.'})
                visuals.append(('P95/P99 thresholds', pd.DataFrame({'metric':['P95','P99'],'value':[p95,p99]})))
                rshare = _share_round_amounts(s_amt)
                if not np.isnan(rshare['p_00']) and rshare['p_00']>params['thr_round']:
                    flags.append({'flag':'High .00 ending share','column':amt,'threshold':params['thr_round'],'value':round(rshare['p_00'],4),'note':'Làm tròn/phát sinh từ nhập tay.'})
                if not np.isnan(rshare['p_50']) and rshare['p_50']>params['thr_round']:
                    flags.append({'flag':'High .50 ending share','column':amt,'threshold':params['thr_round'],'value':round(rshare['p_50'],4),'note':'Pattern giá trị tròn .50 bất thường.'})
                visuals.append(('.00/.50 share', pd.DataFrame([rshare])))
                thrs = _parse_near_thresholds(params.get('near_str',''))
                if thrs:
                    near_tbl = _near_threshold_share(s_amt, thrs, params.get('near_eps_pct',1.0))
                    if not near_tbl.empty:
                        visuals.append(('Near-approval windows', near_tbl))
                        for _,row in near_tbl.iterrows():
                            if row['share']>params['thr_round']:
                                flags.append({'flag':'Near approval threshold cluster','column':amt,'threshold':params['thr_round'],'value':round(float(row['share']),4),
                                              'note': f"Cụm quanh ngưỡng {int(row['threshold']):,} (±{params['near_eps_pct']}%)."})
        dtc = datetime_col if (datetime_col and datetime_col!='(None)' and datetime_col in df.columns) else None
        if dtc:
            t = pd.to_datetime(df[dtc], errors='coerce'); hour = t.dt.hour; weekend = t.dt.dayofweek.isin([5,6])
            if hour.notna().any():
                off_hours = ((hour<7) | (hour>20)).mean()
                if float(off_hours) > params['thr_offh']:
                    flags.append({'flag':'High off‑hours activity','column':dtc,'threshold':params['thr_offh'],'value':round(float(off_hours),4),'note':'Xem lại phân quyền/ca trực/tự động hoá.'})
            if weekend.notna().any():
                w_share = float(weekend.mean())
                if w_share > params['thr_weekend']:
                    flags.append({'flag':'High weekend activity','column':dtc,'threshold':params['thr_weekend'],'value':round(w_share,4),'note':'Rà soát quyền xử lý cuối tuần/quy trình phê duyệt.'})
        if group_cols:
            cols=[c for c in group_cols if c in df.columns]
            if cols:
                ddup = (df[cols].astype(object)
                        .groupby(cols, dropna=False).size().reset_index(name='count').sort_values('count', ascending=False))
                top_dup = ddup[ddup['count'] >= params['dup_min']].head(50)
                if not top_dup.empty:
                    flags.append({'flag':'Duplicate composite keys','column':' + '.join(cols),'threshold':f">={params['dup_min']}",
                                  'value': int(top_dup['count'].max()), 'note':'Rà soát trùng lặp/ghost entries/ghi nhận nhiều lần.'})
                    visuals.append(('Top duplicate keys (≥ threshold)', top_dup))
        if amt and dtc and params.get('use_daily_dups', True):
            tmp = pd.DataFrame({'amt': pd.to_numeric(df[amt], errors='coerce'), 't': pd.to_datetime(df[dtc], errors='coerce')}).dropna()
            if not tmp.empty:
                tmp['date']=tmp['t'].dt.date
                grp_cols = (group_cols or [])
                agg_cols = grp_cols + ['amt','date']
                d2 = tmp.join(df[grp_cols]) if grp_cols else tmp.copy()
                gb = d2.groupby(agg_cols, dropna=False).size().reset_index(name='count').sort_values('count', ascending=False)
                top_amt_dup = gb[gb['count'] >= params['dup_min']].head(50)
                if not top_amt_dup.empty:
                    flags.append({'flag':'Repeated amounts within a day','column':(' + '.join(grp_cols + [amt,'date']) if grp_cols else f'{amt} + date'),
                                  'threshold': f">={params['dup_min']}", 'value': int(top_amt_dup['count'].max()), 'note':'Khả năng chia nhỏ giao dịch / chạy lặp.'})
                    visuals.append(('Same amount duplicates per day', top_amt_dup))
        return flags, visuals

    if run_flags:
        amt_in = None if amount_col=='(None)' else amount_col
        dt_in = None if dt_col=='(None)' else dt_col
        params = dict(thr_zero=thr_zero, thr_tail99=thr_tail99, thr_round=thr_round, thr_offh=thr_offh, thr_weekend=thr_weekend,
                      dup_min=int(dup_min), near_str=near_str, near_eps_pct=near_eps_pct, use_daily_dups=use_daily_dups)
        flags, visuals = compute_fraud_flags(FLAG_DF, amt_in, dt_in, group_cols, params)
        SS['fraud_flags']=flags
        if flags:
            for fl in flags:
                v = to_float(fl.get('value')); thrv = to_float(fl.get('threshold'))
                alarm = '🚨' if (v is not None and thrv is not None and v>thrv) else '🟡'
                st.warning(f"{alarm} [{fl['flag']}] {fl['column']} • thr:{fl.get('threshold')} • val:{fl.get('value')} — {fl['note']}")
        else:
            st.success('🟢 Không có cờ đáng chú ý theo tham số hiện tại.')
        for title, obj in visuals:
            st.markdown(f'**{title}**')
            if isinstance(obj, pd.DataFrame):
                st_df(obj, use_container_width=True, height=min(320, 40+24*min(len(obj),10)))

    with st.expander('🧠 Rule Engine (Flags) — Insights'):
        ctx = build_rule_context(); df_r = evaluate_rules(ctx, scope='flags')
        if not df_r.empty:
            st_df(df_r, use_container_width=True)
        else:
            st.info('Không có rule nào khớp.')
# --------------------------- TAB 7: Risk & Export -----------------------------
with TAB7:
    base_df = DF_FULL
    # ---- Risk summary from Rule Engine v2 (if available) ----
    left, right = st.columns([3,2])
    with left:
        st.subheader('🧭 Automated Risk Assessment — Signals → Next tests → Interpretation')
        # Quick quality & signals (light)
        def _quality_report(df_in: pd.DataFrame) -> tuple[pd.DataFrame, int]:
            rep_rows=[]
            for c in df_in.columns:
                s=df_in[c]
                rep_rows.append({'column':c,'dtype':str(s.dtype),'missing_ratio': round(float(s.isna().mean()),4),
                                 'n_unique':int(s.nunique(dropna=True)),'constant':bool(s.nunique(dropna=True)<=1)})
            dupes=int(df_in.duplicated().sum())
            return pd.DataFrame(rep_rows), dupes
        rep_df, n_dupes = _quality_report(DF_VIEW)
        signals=[]
        if n_dupes>0:
            signals.append({'signal':'Duplicate rows','severity':'Medium','action':'Định nghĩa khoá tổng hợp & walkthrough duplicates'})
        for c in NUM_COLS[:20]:
            s = pd.to_numeric(DF_FULL[c] if SS['df'] is not None else DF_VIEW[c], errors='coerce').replace([np.inf,-np.inf], np.nan).dropna()
            if len(s)==0: continue
            zr=float((s==0).mean()); p99=s.quantile(0.99); share99=float((s>p99).mean())
            if zr>0.30:
                signals.append({'signal':f'Zero‑heavy numeric {c} ({zr:.0%})','severity':'Medium','action':'χ²/Fisher theo đơn vị; review policy/thresholds'})
            if share99>0.02:
                signals.append({'signal':f'Heavy right tail in {c} (>P99 share {share99:.1%})','severity':'High','action':'Benford 1D/2D; cut‑off; outlier review'})
        st_df(pd.DataFrame(signals) if signals else pd.DataFrame([{'status':'No strong risk signals'}]), use_container_width=True, height=320)

    with right:
        st.subheader('🧾 Export (Plotly snapshots) — DOCX / PDF')
        # Figure registry optional — keep minimal by re-capturing on demand in each tab (not stored persistently here)
        st.caption('Chọn nội dung từ các tab, sau đó xuất báo cáo với tiêu đề tuỳ chỉnh.')
        title = st.text_input('Report title', value='Audit Statistics — Findings', key='exp_title')
        scale = st.slider('Export scale (DPI factor)', 1.0, 3.0, 2.0, 0.5, key='exp_scale')
        # For simplicity, take screenshots of figures currently present is not feasible; typical approach is to maintain a registry.
        # Here we export only a simple PDF/DOCX shell with metadata.
        if st.button('🖼️ Export blank shell DOCX/PDF'):
            meta={'title': title, 'file': SS.get('uploaded_name'), 'sha12': SS.get('sha12'), 'time': datetime.now().isoformat(timespec='seconds')}
            docx_path=None; pdf_path=None
            if HAS_DOCX:
                try:
                    d = docx.Document(); d.add_heading(meta['title'], 0)
                    d.add_paragraph(f"File: {meta['file']} • SHA12={meta['sha12']} • Time: {meta['time']}")
                    d.add_paragraph('Gợi ý: quay lại các tab để capture hình (kèm Kaleido) và chèn vào báo cáo.')
                    docx_path = f"report_{int(time.time())}.docx"; d.save(docx_path)
                except Exception: pass
            if HAS_PDF:
                try:
                    doc = fitz.open(); page = doc.new_page(); y=36
                    page.insert_text((36,y), meta['title'], fontsize=16); y+=22
                    page.insert_text((36,y), f"File: {meta['file']} • SHA12={meta['sha12']} • Time: {meta['time']}", fontsize=10); y+=18
                    page.insert_text((36,y), 'Gợi ý: quay lại các tab để capture hình (Kaleido) và chèn vào báo cáo.', fontsize=10)
                    pdf_path = f"report_{int(time.time())}.pdf"; doc.save(pdf_path); doc.close()
                except Exception: pass
            outs=[p for p in [docx_path,pdf_path] if p]
            if outs:
                st.success('Exported: ' + ', '.join(outs))
                for pth in outs:
                    with open(pth,'rb') as f: st.download_button(f'⬇️ Download {os.path.basename(pth)}', data=f.read(), file_name=os.path.basename(pth))
            else:
                st.error('Export failed. Hãy cài python-docx/pymupdf.')

# End of file

